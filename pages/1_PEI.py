# pages/1_PEI.py
import streamlit as st
from datetime import date
from io import BytesIO
from docx import Document
from openai import OpenAI
from pypdf import PdfReader
from fpdf import FPDF
import base64
import json
import os
import time
import re

# ✅ 1) set_page_config (UMA VEZ SÓ e sempre no topo)
st.set_page_config(
    page_title="Omnisfera | PEI",
    page_icon="📘",
    layout="wide",
    initial_sidebar_state="expanded",
)

APP_VERSION = "v150.0 (SaaS Design)"

# ✅ 2) UI lockdown (não quebra se faltar arquivo)
try:
    from ui_lockdown import hide_streamlit_chrome_if_needed, hide_default_sidebar_nav
    hide_streamlit_chrome_if_needed()
    hide_default_sidebar_nav()
except Exception:
    pass

# ✅ 3) Flag de ambiente (opcional)
try:
    IS_TEST_ENV = st.secrets.get("ENV") == "TESTE"
except Exception:
    IS_TEST_ENV = False

# ✅ 4) Gate mínimo: autenticado + workspace_id
if not st.session_state.get("autenticado"):
    st.error("🔒 Acesso negado. Faça login na Página Inicial.")
    st.stop()

ws_id = st.session_state.get("workspace_id")
if not ws_id:
    st.error("Workspace não definido. Volte ao Início e valide o PIN.")
    if st.button("Voltar para Login", key="pei_btn_voltar_login", use_container_width=True):
        for k in ["autenticado", "workspace_id", "workspace_name", "usuario_nome", "usuario_cargo", "supabase_jwt", "supabase_user_id"]:
            st.session_state.pop(k, None)
        st.switch_page("streamlit_app.py")
    st.stop()

# ✅ 5) Supabase (opcional: não bloqueia PEI se der ruim)
sb = None
try:
    from _client import get_supabase
    sb = get_supabase()  # <-- cliente (não é função)
except Exception:
    sb = None

# Guardas legadas (não travam)
def verificar_login_supabase():
    st.session_state.setdefault("supabase_jwt", "")
    st.session_state.setdefault("supabase_user_id", "")

verificar_login_supabase()
OWNER_ID = st.session_state.get("supabase_user_id", "")

# ✅ Sidebar UNIFICADA (navegação + sessão + salvar/carregar + sync)
with st.sidebar:
    st.markdown("### 🧭 Navegação")
    if st.button("🏠 Home", key="pei_nav_home", use_container_width=True):
        st.switch_page("streamlit_app.py")  # se sua home for pages/0_Home.py, troque aqui

    col1, col2 = st.columns(2)
    with col1:
        st.button("📘 PEI", key="pei_nav_pei", use_container_width=True, disabled=True)
    with col2:
        if st.button("🧩 PAEE", key="pei_nav_paee", use_container_width=True):
            st.switch_page("pages/2_PAE.py")

    if st.button("🚀 Hub", key="pei_nav_hub", use_container_width=True):
        st.switch_page("pages/3_Hub_Inclusao.py")

    st.markdown("---")
    st.markdown("### 👤 Sessão")
    st.caption(f"Usuário: **{st.session_state.get('usuario_nome','')}**")
    st.caption(f"Workspace: **{st.session_state.get('workspace_name','')}**")

    st.markdown("---")
    st.markdown("### 🔑 OpenAI")
    if 'OPENAI_API_KEY' in st.secrets:
        api_key = st.secrets['OPENAI_API_KEY']
        st.success("✅ OpenAI OK")
    else:
        api_key = st.text_input("Chave OpenAI:", type="password", key="pei_openai_key")

    st.markdown("---")
    st.markdown("### 🧾 Status do Aluno (Supabase)")
    st.session_state.setdefault("selected_student_id", None)
    st.session_state.setdefault("selected_student_name", "")

    student_id = st.session_state.get("selected_student_id")
    if student_id:
        st.success("✅ Vinculado ao Supabase")
        st.caption(f"student_id: {student_id[:8]}...")
    else:
        st.warning("📝 Rascunho (ainda não salvo no Supabase)")

    # Aviso se supabase não estiver pronto
    if sb is None:
        st.info("Supabase não inicializado (sb=None). O PEI funciona em rascunho, mas não salva/carrega.")

    st.markdown("---")


# ==============================================================================
# 0. CONFIGURAÇÃO DE PÁGINA
# ==============================================================================
st.set_page_config(
    page_title="Omnisfera | PEI",
    page_icon="📘",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==============================================================================
# 1. GUARDAS (LOGIN + SUPABASE)
# ==============================================================================
def verificar_login_app():
    if "autenticado" not in st.session_state or not st.session_state["autenticado"]:
        st.error("🔒 Acesso Negado. Faça login na Página Inicial.")
        st.stop()

def verificar_login_supabase():
    # Supabase é necessário para SALVAR/CARREGAR, mas o PEI pode abrir como rascunho.
    # Então aqui só avisamos, não bloqueamos tudo.
    if "supabase_jwt" not in st.session_state or not st.session_state["supabase_jwt"]:
        st.session_state["supabase_jwt"] = ""
    if "supabase_user_id" not in st.session_state or not st.session_state["supabase_user_id"]:
        st.session_state["supabase_user_id"] = ""

verificar_login_app()
verificar_login_supabase()

def sb():
    return get_supabase_user(st.session_state["supabase_jwt"])

OWNER_ID = st.session_state.get("supabase_user_id", "")

# ==============================================================================
# 2. SUPABASE: STUDENTS (criar/atualizar/listar/excluir) — com workspace_id
# ==============================================================================
def _sb_ok():
    return sb is not None and OWNER_ID and ws_id

def db_create_student(payload: dict):
    if not _sb_ok():
        raise RuntimeError("Supabase não está pronto (sb/OWNER_ID/workspace_id).")
    payload = dict(payload or {})
    payload["owner_id"] = OWNER_ID
    payload["workspace_id"] = ws_id
    res = sb.table("students").insert(payload).execute()
    return (res.data or [None])[0]

def db_update_student(student_id: str, payload: dict):
    if not _sb_ok():
        raise RuntimeError("Supabase não está pronto (sb/OWNER_ID/workspace_id).")
    payload = dict(payload or {})
    sb.table("students").update(payload).eq("id", student_id).eq("owner_id", OWNER_ID).eq("workspace_id", ws_id).execute()

def db_delete_student(student_id: str):
    if not _sb_ok():
        raise RuntimeError("Supabase não está pronto (sb/OWNER_ID/workspace_id).")
    sb.table("students").delete().eq("id", student_id).eq("owner_id", OWNER_ID).eq("workspace_id", ws_id).execute()

def db_list_students(search: str | None = None):
    if not _sb_ok():
        return []
    q = (
        sb.table("students")
        .select("id, owner_id, workspace_id, name, birth_date, grade, class_group, diagnosis, created_at, updated_at")
        .eq("owner_id", OWNER_ID)
        .eq("workspace_id", ws_id)
        .order("name", desc=False)
    )
    if search:
        q = q.ilike("name", f"%{search}%")
    res = q.execute()
    return res.data or []

def db_get_student(student_id: str):
    if not _sb_ok():
        return None
    res = (
        sb.table("students")
        .select("id, owner_id, workspace_id, name, birth_date, grade, class_group, diagnosis, created_at, updated_at")
        .eq("id", student_id)
        .eq("owner_id", OWNER_ID)
        .eq("workspace_id", ws_id)
        .limit(1)
        .execute()
    )
    data = res.data or []
    return data[0] if data else None


# ==============================================================================
# 3. BLOCO VISUAL (badge / logo)
# ==============================================================================
def get_logo_base64():
    caminhos = ["omni_icone.png", "logo.png", "iconeaba.png", "omni.png", "ominisfera.png"]
    for c in caminhos:
        if os.path.exists(c):
            with open(c, "rb") as f:
                return f"data:image/png;base64,{base64.b64encode(f.read()).decode()}"
    # fallback
    return "https://cdn-icons-png.flaticon.com/512/1183/1183672.png"

src_logo_giratoria = get_logo_base64()

if IS_TEST_ENV:
    card_bg = "rgba(255, 220, 50, 0.95)"
    card_border = "rgba(200, 160, 0, 0.5)"
else:
    card_bg = "rgba(255, 255, 255, 0.85)"
    card_border = "rgba(255, 255, 255, 0.6)"

st.markdown(f"""
<style>
    .omni-badge {{
        position: fixed; top: 15px; right: 15px;
        background: {card_bg}; border: 1px solid {card_border};
        backdrop-filter: blur(8px); -webkit-backdrop-filter: blur(8px);
        padding: 4px 30px; min-width: 260px; justify-content: center;
        border-radius: 20px; box-shadow: 0 4px 15px rgba(0,0,0,0.08);
        z-index: 999990; display: flex; align-items: center; gap: 10px;
        pointer-events: none;
    }}
    .omni-text {{
        font-family: 'Nunito', sans-serif; font-weight: 800; font-size: 0.9rem;
        color: #2D3748; letter-spacing: 1px; text-transform: uppercase;
    }}
    @keyframes spin-slow {{ from {{ transform: rotate(0deg); }} to {{ transform: rotate(360deg); }} }}
    .omni-logo-spin {{ height: 26px; width: 26px; animation: spin-slow 10s linear infinite; }}
</style>
<div class="omni-badge">
    <img src="{src_logo_giratoria}" class="omni-logo-spin">
    <span class="omni-text">OMNISFERA</span>
</div>
""", unsafe_allow_html=True)


# ==============================================================================
# 4. LISTAS DE DADOS
# ==============================================================================
LISTA_SERIES = [
    "Educação Infantil (Creche)", "Educação Infantil (Pré-Escola)",
    "1º Ano (Fund. I)", "2º Ano (Fund. I)", "3º Ano (Fund. I)", "4º Ano (Fund. I)", "5º Ano (Fund. I)",
    "6º Ano (Fund. II)", "7º Ano (Fund. II)", "8º Ano (Fund. II)", "9º Ano (Fund. II)",
    "1ª Série (EM)", "2ª Série (EM)", "3ª Série (EM)", "EJA (Educação de Jovens e Adultos)"
]

LISTA_ALFABETIZACAO = [
    "Não se aplica (Educação Infantil)",
    "Pré-Silábico (Garatuja/Desenho sem letras)",
    "Pré-Silábico (Letras aleatórias sem valor sonoro)",
    "Silábico (Sem valor sonoro convencional)",
    "Silábico (Com valor sonoro vogais/consoantes)",
    "Silábico-Alfabético (Transição)",
    "Alfabético (Escrita fonética, com erros ortográficos)",
    "Ortográfico (Escrita convencional consolidada)"
]

LISTAS_BARREIRAS = {
    "Funções Cognitivas": ["Atenção Sustentada/Focada", "Memória de Trabalho (Operacional)", "Flexibilidade Mental", "Planejamento e Organização", "Velocidade de Processamento", "Abstração e Generalização"],
    "Comunicação e Linguagem": ["Linguagem Expressiva (Fala)", "Linguagem Receptiva (Compreensão)", "Pragmática (Uso social da língua)", "Processamento Auditivo", "Intenção Comunicativa"],
    "Socioemocional": ["Regulação Emocional (Autocontrole)", "Tolerância à Frustração", "Interação Social com Pares", "Autoestima e Autoimagem", "Reconhecimento de Emoções"],
    "Sensorial e Motor": ["Praxias Globais (Coordenação Grossa)", "Praxias Finas (Coordenação Fina)", "Hipersensibilidade Sensorial", "Hipossensibilidade (Busca Sensorial)", "Planejamento Motor"],
    "Acadêmico": ["Decodificação Leitora", "Compreensão Textual", "Raciocínio Lógico-Matemático", "Grafomotricidade (Escrita manual)", "Produção Textual"]
}

LISTA_POTENCIAS = ["Memória Visual", "Musicalidade/Ritmo", "Interesse em Tecnologia", "Hiperfoco Construtivo", "Liderança Natural", "Habilidades Cinestésicas (Esportes)", "Expressão Artística (Desenho)", "Cálculo Mental Rápido", "Oralidade/Vocabulário", "Criatividade/Imaginação", "Empatia/Cuidado com o outro", "Resolução de Problemas", "Curiosidade Investigativa"]

LISTA_PROFISSIONAIS = ["Psicólogo Clínico", "Neuropsicólogo", "Fonoaudiólogo", "Terapeuta Ocupacional", "Neuropediatra", "Psiquiatra Infantil", "Psicopedagogo Clínico", "Professor de Apoio (Mediador)", "Acompanhante Terapêutico (AT)", "Musicoterapeuta", "Equoterapeuta", "Oftalmologista"]

LISTA_FAMILIA = ["Mãe", "Mãe 2", "Pai", "Pai 2", "Madrasta", "Padrasto", "Avó Materna", "Avó Paterna", "Avô Materno", "Avô Paterno", "Irmãos", "Tios", "Primos", "Tutor Legal", "Abrigo Institucional"]


# ==============================================================================
# 5. ESTADO DEFAULT (RASCUNHO)
# ==============================================================================
default_state = {
    "nome": "",
    "nasc": date(2015, 1, 1),
    "serie": None,
    "turma": "",
    "diagnostico": "",
    "lista_medicamentos": [],
    "composicao_familiar_tags": [],
    "historico": "",
    "familia": "",
    "hiperfoco": "",
    "potencias": [],
    "rede_apoio": [],
    "orientacoes_especialistas": "",
    "orientacoes_por_profissional": {},
    "checklist_evidencias": {},
    "nivel_alfabetizacao": "Não se aplica (Educação Infantil)",
    "barreiras_selecionadas": {k: [] for k in LISTAS_BARREIRAS.keys()},
    "niveis_suporte": {},
    "observacoes_barreiras": {},
    "estrategias_acesso": [],
    "estrategias_ensino": [],
    "estrategias_avaliacao": [],
    "ia_sugestao": "",
    "ia_mapa_texto": "",
    "outros_acesso": "",
    "outros_ensino": "",
    "monitoramento_data": date.today(),
    "status_meta": "Não Iniciado",
    "parecer_geral": "Manter Estratégias",
    "proximos_passos_select": [],
    "status_validacao_pei": "rascunho",
    "feedback_ajuste": "",
    "status_validacao_game": "rascunho",
    "feedback_ajuste_game": "",
    "matricula": "",
    "meds_extraidas_tmp": [],
    "status_meds_extraidas": "idle",
}

if "dados" not in st.session_state:
    st.session_state.dados = default_state
else:
    for k, v in default_state.items():
        if k not in st.session_state.dados:
            st.session_state.dados[k] = v

st.session_state.setdefault("pdf_text", "")

# vínculo supabase
st.session_state.setdefault("selected_student_id", None)
st.session_state.setdefault("selected_student_name", "")


# ==============================================================================
# 6. SUPABASE: carregar/salvar PEI (pei_documents) — só quando vinculado
# ==============================================================================
def supa_load_latest_pei(student_id: str):
    if not _sb_ok():
        return None
    res = (
        sb.table("pei_documents")
        .select("*")
        .eq("student_id", student_id)
        .eq("owner_id", OWNER_ID)
        .eq("workspace_id", ws_id)
        .order("updated_at", desc=True)
        .limit(1)
        .execute()
    )
    data = res.data or []
    return data[0] if data else None

def supa_save_pei(student_id: str, payload: dict, pdf_text: str):
    if not _sb_ok():
        raise RuntimeError("Supabase não está pronto (sb/OWNER_ID/workspace_id).")

    def _jsonify(x):
        return json.loads(json.dumps(x, default=str))

    safe_payload = _jsonify(payload or {})
    year = date.today().year

    existing = supa_load_latest_pei(student_id)
    if existing:
        sb.table("pei_documents").update({
            "payload": safe_payload,
            "pdf_text": (pdf_text or "")[:20000],
            "school_year": year,
            "status": (payload or {}).get("status_validacao_pei", "draft"),
        }).eq("id", existing["id"]).eq("owner_id", OWNER_ID).eq("workspace_id", ws_id).execute()
    else:
        sb.table("pei_documents").insert({
            "owner_id": OWNER_ID,
            "workspace_id": ws_id,
            "student_id": student_id,
            "school_year": year,
            "status": (payload or {}).get("status_validacao_pei", "draft"),
            "payload": safe_payload,
            "pdf_text": (pdf_text or "")[:20000],
        }).execute()

def supa_sync_student_from_dados(student_id: str, d: dict):
    # mantém students atualizado com dados básicos do PEI
    db_update_student(student_id, {
        "name": d.get("nome") or None,
        "birth_date": d.get("nasc").isoformat() if hasattr(d.get("nasc"), "isoformat") else None,
        "grade": d.get("serie") or None,
        "class_group": d.get("turma") or None,
        "diagnosis": d.get("diagnostico") or None,
    })


# ==============================================================================
# 7. UTILITÁRIOS
# ==============================================================================
def calcular_progresso():
    if st.session_state.dados.get('ia_sugestao'):
        return 100
    pontos = 0
    total = 7
    d = st.session_state.dados
    if d.get('nome'):
        pontos += 1
    if d.get('serie'):
        pontos += 1
    if d.get('nivel_alfabetizacao') and d.get('nivel_alfabetizacao') != 'Não se aplica (Educação Infantil)':
        pontos += 1
    if any(d.get('checklist_evidencias', {}).values()):
        pontos += 1
    if d.get('hiperfoco'):
        pontos += 1
    if any(d.get('barreiras_selecionadas', {}).values()):
        pontos += 1
    if d.get('estrategias_ensino'):
        pontos += 1
    return int((pontos / total) * 90)

def finding_logo():
    possiveis = ["360.png", "360.jpg", "logo.png", "logo.jpg", "iconeaba.png"]
    for nome in possiveis:
        if os.path.exists(nome):
            return nome
    return None

def get_base64_image(image_path):
    if not image_path or not os.path.exists(image_path):
        return ""
    with open(image_path, "rb") as img_file:
        return base64.b64encode(img_file.read()).decode()

def ler_pdf(arquivo):
    try:
        reader = PdfReader(arquivo)
        texto = ""
        for i, page in enumerate(reader.pages):
            if i >= 6:
                break
            texto += (page.extract_text() or "") + "\n"
        return texto
    except:
        return ""

def render_progresso():
    p = calcular_progresso()
    icon_html = f'<img src="{src_logo_giratoria}" class="omni-logo-spin" style="width: 25px; height: 25px;">'
    bar_color = "linear-gradient(90deg, #FF6B6B 0%, #FF8E53 100%)"
    if p >= 100:
        bar_color = "linear-gradient(90deg, #00C6FF 0%, #0072FF 100%)"
    st.markdown(
        f"""<div style="width:100%; margin: 0 0 20px 0;">
              <div style="width:100%; height:3px; background:#E2E8F0; border-radius:2px; position:relative;">
                <div style="height:3px; width:{p}%; background:{bar_color}; border-radius:2px;"></div>
                <div style="position:absolute; top:-14px; left:{p}%; transform:translateX(-50%);">{icon_html}</div>
              </div>
            </div>""",
        unsafe_allow_html=True
    )
# ==============================================================================
# 7B. UTILITÁRIOS AVANÇADOS (idade, segmento, metas, radar, etc.)
# ==============================================================================

def calcular_idade(data_nasc):
    if not data_nasc:
        return ""
    hoje = date.today()
    idade = hoje.year - data_nasc.year - ((hoje.month, hoje.day) < (data_nasc.month, data_nasc.day))
    return f"{idade} anos"

def detectar_nivel_ensino(serie_str: str | None):
    if not serie_str:
        return "INDEFINIDO"
    s = serie_str.lower()
    if "infantil" in s:
        return "EI"
    if "1º ano" in s or "2º ano" in s or "3º ano" in s or "4º ano" in s or "5º ano" in s:
        return "FI"
    if "6º ano" in s or "7º ano" in s or "8º ano" in s or "9º ano" in s:
        return "FII"
    if "série" in s or "médio" in s or "eja" in s:
        return "EM"
    return "INDEFINIDO"

def get_segmento_info_visual(serie: str | None):
    nivel = detectar_nivel_ensino(serie or "")
    if nivel == "EI":
        return "Educação Infantil", "#4299e1", "Foco: Campos de Experiência (BNCC)."
    if nivel == "FI":
        return "Anos Iniciais (Fund. I)", "#48bb78", "Foco: Alfabetização e BNCC."
    if nivel == "FII":
        return "Anos Finais (Fund. II)", "#ed8936", "Foco: Autonomia e Identidade."
    if nivel == "EM":
        return "Ensino Médio / EJA", "#9f7aea", "Foco: Projeto de Vida."
    return "Selecione a Série", "grey", "Aguardando seleção..."

def get_hiperfoco_emoji(texto: str | None):
    if not texto:
        return "🚀"
    t = texto.lower()
    if "jogo" in t or "game" in t or "minecraft" in t or "roblox" in t:
        return "🎮"
    if "dino" in t:
        return "🦖"
    if "fute" in t or "bola" in t:
        return "⚽"
    if "desenho" in t or "arte" in t:
        return "🎨"
    if "músic" in t or "music" in t:
        return "🎵"
    if "anim" in t or "gato" in t or "cachorro" in t:
        return "🐾"
    if "carro" in t:
        return "🏎️"
    if "espaço" in t or "espaco" in t:
        return "🪐"
    return "🚀"

def calcular_complexidade_pei(dados: dict):
    n_bar = sum(len(v) for v in (dados.get("barreiras_selecionadas") or {}).values())
    n_suporte_alto = sum(
        1 for v in (dados.get("niveis_suporte") or {}).values()
        if v in ["Substancial", "Muito Substancial"]
    )
    recursos = 0
    if dados.get("rede_apoio"):
        recursos += 3
    if dados.get("lista_medicamentos"):
        recursos += 2
    saldo = (n_bar + n_suporte_alto) - recursos
    if saldo <= 2:
        return "FLUIDA", "#F0FFF4", "#276749"
    if saldo <= 7:
        return "ATENÇÃO", "#FFFFF0", "#D69E2E"
    return "CRÍTICA", "#FFF5F5", "#C53030"

def extrair_tag_ia(texto: str, tag: str):
    if not texto:
        return ""
    padrao = fr"\[{tag}\](.*?)(\[|$)"
    match = re.search(padrao, texto, re.DOTALL)
    return match.group(1).strip() if match else ""

def extrair_metas_estruturadas(texto: str):
    bloco = extrair_tag_ia(texto or "", "METAS_SMART")
    metas = {"Curto": "Definir...", "Medio": "Definir...", "Longo": "Definir..."}
    if bloco:
        linhas = bloco.split("\n")
        for l in linhas:
            l_clean = re.sub(r"^[\-\*]+", "", l).strip()
            if not l_clean:
                continue
            if "Curto" in l or "2 meses" in l:
                metas["Curto"] = l_clean.split(":")[-1].strip()
            elif "Médio" in l or "Semestre" in l or "Medio" in l:
                metas["Medio"] = l_clean.split(":")[-1].strip()
            elif "Longo" in l or "Ano" in l:
                metas["Longo"] = l_clean.split(":")[-1].strip()
    return metas

def get_pro_icon(nome_profissional: str | None):
    p = (nome_profissional or "").lower()
    if "psic" in p:
        return "🧠"
    if "fono" in p:
        return "🗣️"
    if "terapeuta" in p or "equo" in p or "musico" in p:
        return "🧩"
    if "neuro" in p or "psiq" in p or "medico" in p:
        return "🩺"
    return "👨‍⚕️"

def inferir_componentes_impactados(dados: dict):
    barreiras = dados.get("barreiras_selecionadas", {}) or {}
    serie = (dados.get("serie") or "")
    nivel = detectar_nivel_ensino(serie)
    impactados = set()

    # Leitura
    if barreiras.get("Acadêmico") and any("Leitora" in b for b in barreiras["Acadêmico"]):
        impactados.add("Língua Portuguesa")
        impactados.add("História/Sociologia/Filosofia" if nivel == "EM" else "História/Geografia")

    # Matemática
    if barreiras.get("Acadêmico") and any("Matemático" in b for b in barreiras["Acadêmico"]):
        impactados.add("Matemática")
        if nivel == "EM":
            impactados.add("Física/Química")
        elif nivel == "FII":
            impactados.add("Ciências")

    # Cognitivas (transversal)
    if barreiras.get("Funções Cognitivas"):
        impactados.add("Transversal (Todas as áreas)")

    # Motor fino
    if barreiras.get("Sensorial e Motor") and any("Fina" in b for b in barreiras["Sensorial e Motor"]):
        impactados.add("Arte")
        impactados.add("Geometria")

    if not impactados and dados.get("diagnostico"):
        return ["Análise Geral (Baseada no Diagnóstico)"]

    return list(impactados) if impactados else ["Nenhum componente específico detectado automaticamente"]


# ==============================================================================
# 7C. PDF / DOCX (Exportação)
# ==============================================================================

def limpar_texto_pdf(texto: str):
    if not texto:
        return ""
    t = texto.replace("**", "").replace("__", "").replace("#", "").replace("•", "-")
    t = t.replace("“", '"').replace("”", '"').replace("‘", "'").replace("’", "'")
    t = t.replace("–", "-").replace("—", "-")
    return t.encode("latin-1", "replace").decode("latin-1")

class PDF_Classic(FPDF):
    def header(self):
        self.set_fill_color(248, 248, 248)
        self.rect(0, 0, 210, 40, "F")
        logo = finding_logo()
        x_offset = 40 if logo else 12
        if logo:
            self.image(logo, 10, 8, 25)
        self.set_xy(x_offset, 12)
        self.set_font("Arial", "B", 14)
        self.set_text_color(50, 50, 50)
        self.cell(0, 8, "PEI - PLANO DE ENSINO INDIVIDUALIZADO", 0, 1, "L")
        self.set_xy(x_offset, 19)
        self.set_font("Arial", "", 9)
        self.set_text_color(100, 100, 100)
        self.cell(0, 5, "Documento de Planejamento e Flexibilização Curricular", 0, 1, "L")
        self.ln(15)

    def footer(self):
        self.set_y(-15)
        self.set_font("Arial", "I", 8)
        self.set_text_color(150, 150, 150)
        self.cell(0, 10, f"Página {self.page_no()} | Gerado via Omnisfera", 0, 0, "C")

    def section_title(self, label):
        self.ln(6)
        self.set_fill_color(230, 230, 230)
        self.rect(10, self.get_y(), 190, 8, "F")
        self.set_font("ZapfDingbats", "", 10)
        self.set_text_color(80, 80, 80)
        self.set_xy(12, self.get_y() + 1)
        self.cell(5, 6, "o", 0, 0)
        self.set_font("Arial", "B", 11)
        self.set_text_color(50, 50, 50)
        self.cell(0, 6, label.upper(), 0, 1, "L")
        self.ln(4)

    def add_flat_icon_item(self, texto, bullet_type="check"):
        self.set_font("ZapfDingbats", "", 10)
        self.set_text_color(80, 80, 80)
        char = "3" if bullet_type == "check" else "l"
        self.cell(6, 5, char, 0, 0)
        self.set_font("Arial", "", 10)
        self.set_text_color(0)
        self.multi_cell(0, 5, texto)
        self.ln(1)

class PDF_Simple_Text(FPDF):
    def header(self):
        self.set_font("Arial", "B", 16)
        self.set_text_color(50)
        self.cell(0, 10, "ROTEIRO DE MISSÃO", 0, 1, "C")
        self.set_draw_color(150)
        self.line(10, 25, 200, 25)
        self.ln(10)

def gerar_pdf_final(dados: dict):
    pdf = PDF_Classic()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)

    pdf.section_title("Identificação e Contexto")
    pdf.set_font("Arial", "B", 10)
    pdf.cell(35, 6, "Estudante:", 0, 0)
    pdf.set_font("Arial", "", 10)
    pdf.cell(0, 6, dados.get("nome", ""), 0, 1)

    pdf.set_font("Arial", "B", 10)
    pdf.cell(35, 6, "Série/Turma:", 0, 0)
    pdf.set_font("Arial", "", 10)
    pdf.cell(0, 6, f"{dados.get('serie','')} - {dados.get('turma','')}", 0, 1)

    pdf.set_font("Arial", "B", 10)
    pdf.cell(35, 6, "Diagnóstico:", 0, 0)
    pdf.set_font("Arial", "", 10)
    pdf.multi_cell(0, 6, dados.get("diagnostico", ""))
    pdf.ln(2)

    if any((dados.get("barreiras_selecionadas") or {}).values()):
        pdf.section_title("Plano de Suporte (Barreiras x Nível)")
        for area, itens in (dados.get("barreiras_selecionadas") or {}).items():
            if itens:
                pdf.set_font("Arial", "B", 10)
                pdf.cell(0, 8, limpar_texto_pdf(area), 0, 1)
                for item in itens:
                    nivel = (dados.get("niveis_suporte") or {}).get(f"{area}_{item}", "Monitorado")
                    pdf.add_flat_icon_item(limpar_texto_pdf(f"{item} (Nível: {nivel})"), "check")

    if dados.get("ia_sugestao"):
        pdf.add_page()
        pdf.section_title("Planejamento Pedagógico Detalhado")
        texto_limpo = limpar_texto_pdf(dados["ia_sugestao"])
        texto_limpo = re.sub(r"\[.*?\]", "", texto_limpo)

        for linha in texto_limpo.split("\n"):
            l = linha.strip()
            if not l:
                continue
            if l.startswith("###") or l.startswith("##"):
                pdf.ln(5)
                pdf.set_font("Arial", "B", 12)
                pdf.set_text_color(0, 51, 102)
                pdf.cell(0, 8, l.replace("#", "").strip(), 0, 1, "L")
                pdf.set_font("Arial", "", 10)
                pdf.set_text_color(0, 0, 0)
            elif l.startswith("-") or l.startswith("*"):
                pdf.add_flat_icon_item(l.replace("-", "").replace("*", "").strip(), "dot")
            else:
                pdf.multi_cell(0, 6, l)

    return pdf.output(dest="S").encode("latin-1", "replace")

def gerar_pdf_tabuleiro_simples(texto: str):
    pdf = PDF_Simple_Text()
    pdf.add_page()
    pdf.set_font("Arial", size=11)
    for linha in limpar_texto_pdf(texto).split("\n"):
        l = linha.strip()
        if not l:
            continue
        if l.isupper() or "**" in linha:
            pdf.ln(4)
            pdf.set_font("Arial", "B", 11)
            pdf.set_fill_color(240, 240, 240)
            pdf.cell(0, 8, l.replace("**", ""), 0, 1, "L", fill=True)
            pdf.set_font("Arial", "", 11)
        else:
            pdf.multi_cell(0, 6, l)
    return pdf.output(dest="S").encode("latin-1", "ignore")

def gerar_docx_final(dados: dict):
    doc = Document()
    doc.add_heading("PEI - " + (dados.get("nome") or "Sem Nome"), 0)
    if dados.get("ia_sugestao"):
        doc.add_paragraph(re.sub(r"\[.*?\]", "", dados["ia_sugestao"]))
    b = BytesIO()
    doc.save(b)
    b.seek(0)
    return b


# ==============================================================================
# 7D. IA (Extração PDF + Consultoria + Gamificação)
# ==============================================================================

def extrair_dados_pdf_ia(api_key: str, texto_pdf: str):
    if not api_key:
        return None, "Configure a Chave API OpenAI."
    try:
        client = OpenAI(api_key=api_key)
        prompt = (
            "Analise este laudo médico/escolar. Extraia: 1) Diagnóstico; 2) Medicamentos. "
            'Responda em JSON no formato: { "diagnostico": "...", "medicamentos": [ {"nome": "...", "posologia": "..."} ] }. '
            f"Texto: {texto_pdf[:4000]}"
        )
        res = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"},
        )
        return json.loads(res.choices[0].message.content), None
    except Exception as e:
        return None, str(e)

def consultar_gpt_pedagogico(api_key: str, dados: dict, contexto_pdf: str = "", modo_pratico: bool = False, feedback_usuario: str = ""):
    if not api_key:
        return None, "⚠️ Configure a Chave API OpenAI."
    try:
        client = OpenAI(api_key=api_key)

        evid = "\n".join([f"- {k.replace('?', '')}" for k, v in (dados.get("checklist_evidencias") or {}).items() if v])
        meds_info = "\n".join(
            [f"- {m.get('nome','')} ({m.get('posologia','')})." for m in (dados.get("lista_medicamentos") or [])]
        ) if dados.get("lista_medicamentos") else "Nenhuma medicação informada."

        hiperfoco_txt = f"HIPERFOCO DO ALUNO: {dados.get('hiperfoco','')}" if dados.get("hiperfoco") else "Hiperfoco: Não identificado."
        serie = dados.get("serie") or ""
        nivel_ensino = detectar_nivel_ensino(serie)
        alfabetizacao = dados.get("nivel_alfabetizacao", "Não Avaliado")

        prompt_identidade = f"""
[PERFIL_NARRATIVO]
Inicie com "👤 QUEM É O ESTUDANTE?". Crie um parágrafo humanizado. {hiperfoco_txt}.
Use o hiperfoco para conectar com a aprendizagem.
[/PERFIL_NARRATIVO]
""".strip()

        prompt_diagnostico = """
### 1. 🏥 DIAGNÓSTICO E IMPACTO (FUNDAMENTAL):
- Cite o Diagnóstico (e o CID se disponível).
- Descreva os impactos diretos na aprendizagem.
- Liste cuidados/pontos de atenção.
""".strip()

        prompt_literacia = ""
        if "Alfabético" not in alfabetizacao and alfabetizacao != "Não se aplica (Educação Infantil)":
            prompt_literacia = f"""[ATENÇÃO CRÍTICA: ALFABETIZAÇÃO] Fase: {alfabetizacao}. Inclua 2 ações de consciência fonológica.[/ATENÇÃO CRÍTICA]"""

        prompt_hub = """
### 6. 🧩 CHECKLIST DE ADAPTAÇÃO E ACESSIBILIDADE:
**A. Mediação (Triângulo de Ouro):**
1) Instruções passo a passo
2) Fragmentação de tarefas
3) Scaffolding

**B. Acessibilidade:**
4) Inferências/figuras de linguagem
5) Descrição de imagens (alt text)
6) Adaptação visual (fonte/espaçamento)
7) Adequação de desafio
""".strip()

        prompt_componentes = ""
        if nivel_ensino != "EI":
            prompt_componentes = f"""
### 4. ⚠️ COMPONENTES CURRICULARES DE ATENÇÃO:
Com base no diagnóstico ({dados.get('diagnostico','')}) e nas barreiras citadas, identifique componentes que exigirão maior flexibilização.
- Liste componentes
- Para cada um, explique o motivo técnico
""".strip()

        prompt_metas = """
[METAS_SMART]
- Meta de Curto Prazo (2 meses): [Descreva a meta]
- Meta de Médio Prazo (1 semestre): [Descreva a meta]
- Meta de Longo Prazo (1 ano): [Descreva a meta]
[/METAS_SMART]
""".strip()

        if nivel_ensino == "EI":
            perfil_ia = "Especialista em EDUCAÇÃO INFANTIL e BNCC."
            estrutura_req = f"""
{prompt_identidade}
{prompt_diagnostico}

### 2. 🌟 AVALIAÇÃO DE REPERTÓRIO:
[CAMPOS_EXPERIENCIA_PRIORITARIOS] Destaque 2 ou 3 Campos BNCC. [/CAMPOS_EXPERIENCIA_PRIORITARIOS]

### 3. 🚀 ESTRATÉGIAS DE INTERVENÇÃO:
(Estratégias de acolhimento, rotina e adaptação sensorial).

{prompt_metas}

### 5. ⚠️ PONTOS DE ATENÇÃO FARMACOLÓGICA:
[ANALISE_FARMA] Se houver medicação, cite efeitos colaterais para atenção pedagógica. [/ANALISE_FARMA]

{prompt_hub}
""".strip()
        else:
            perfil_ia = "Especialista em Inclusão Escolar e BNCC."
            instrucao_bncc = "[MAPEAMENTO_BNCC] Separe por Componente Curricular. Inclua código alfanumérico (ex: EF01LP02). [/MAPEAMENTO_BNCC]"
            instrucao_bloom = "[TAXONOMIA_BLOOM] Explique a categoria cognitiva escolhida. [/TAXONOMIA_BLOOM]"
            estrutura_req = f"""
{prompt_identidade}
{prompt_diagnostico}

### 2. 🌟 AVALIAÇÃO DE REPERTÓRIO:
- Defasagens (anos anteriores)
- Foco do ano atual
{instrucao_bncc}
{instrucao_bloom}

### 3. 🚀 ESTRATÉGIAS DE INTERVENÇÃO:
(Adaptações curriculares e de acesso).
{prompt_literacia}

{prompt_componentes}

{prompt_metas}

### 5. ⚠️ PONTOS DE ATENÇÃO FARMACOLÓGICA:
[ANALISE_FARMA] Se houver medicação, cite efeitos colaterais para atenção pedagógica. [/ANALISE_FARMA]

{prompt_hub}
""".strip()

        prompt_feedback = f"AJUSTE SOLICITADO: {feedback_usuario}" if feedback_usuario else ""
        prompt_formatacao = "IMPORTANTE: Use Markdown simples. Use títulos H3 (###). Evite tabelas."

        prompt_sys = f"""{perfil_ia}
MISSÃO: Criar PEI Técnico Oficial.
ESTRUTURA OBRIGATÓRIA:
{estrutura_req}

{prompt_feedback}
{prompt_formatacao}
""".strip()

        if modo_pratico:
            prompt_sys = f"""{perfil_ia}
GUIA PRÁTICO PARA SALA DE AULA.
{prompt_feedback}

{prompt_hub}
""".strip()

        prompt_user = (
            f"ALUNO: {dados.get('nome','')} | SÉRIE: {serie} | HISTÓRICO: {dados.get('historico','')} | "
            f"DIAGNÓSTICO: {dados.get('diagnostico','')} | MEDS: {meds_info} | "
            f"EVIDÊNCIAS: {evid} | LAUDO: {(contexto_pdf or '')[:3000]}"
        )

        res = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "system", "content": prompt_sys}, {"role": "user", "content": prompt_user}],
        )
        return res.choices[0].message.content, None

    except Exception as e:
        return None, str(e)

def gerar_roteiro_gamificado(api_key: str, dados: dict, pei_tecnico: str, feedback_game: str = ""):
    if not api_key:
        return None, "Configure a chave OpenAI."
    try:
        client = OpenAI(api_key=api_key)
        serie = dados.get("serie") or ""
        nivel_ensino = detectar_nivel_ensino(serie)
        hiperfoco = dados.get("hiperfoco") or "brincadeiras"
        nome_curto = (dados.get("nome","").split() or ["Estudante"])[0]

        contexto_seguro = (
            f"ALUNO: {nome_curto} | HIPERFOCO: {hiperfoco} | "
            f"PONTOS FORTES: {', '.join(dados.get('potencias',[]))}"
        )

        prompt_feedback = f"AJUSTE: {feedback_game}" if feedback_game else ""

        if nivel_ensino == "EI":
            prompt_sys = "Crie uma história visual (4-5 anos) com emojis. Estrutura: começo, desafio, ajuda, conquista, rotina."
        elif nivel_ensino == "FI":
            prompt_sys = "Crie um quadro de missões RPG (6-10 anos). Estrutura: mapa, missões, recompensas, superpoder."
        else:
            prompt_sys = "Crie uma ficha RPG (adolescente). Estrutura: quest, skills, buffs, checklists e metas."

        full_sys = f"{prompt_sys}\n{prompt_feedback}"
        res = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "system", "content": full_sys}, {"role": "user", "content": contexto_seguro}],
        )
        return res.choices[0].message.content, None
    except Exception as e:
        return None, str(e)


# ==============================================================================
# 7E. AÇÕES AUXILIARES (Reset)
# ==============================================================================
def limpar_formulario():
    # recria um "rascunho limpo" preservando a estrutura do dicionário
    st.session_state.dados = {
        'nome': '',
        'nasc': date(2015, 1, 1),
        'serie': None,
        'turma': '',
        'diagnostico': '',
        'lista_medicamentos': [],
        'composicao_familiar_tags': [],
        'historico': '',
        'familia': '',
        'hiperfoco': '',
        'potencias': [],
        'rede_apoio': [],
        'orientacoes_especialistas': '',
        'checklist_evidencias': {},
        'nivel_alfabetizacao': 'Não se aplica (Educação Infantil)',
        'barreiras_selecionadas': {k: [] for k in LISTAS_BARREIRAS.keys()},
        'niveis_suporte': {},
        'estrategias_acesso': [],
        'estrategias_ensino': [],
        'estrategias_avaliacao': [],
        'ia_sugestao': '',
        'ia_mapa_texto': '',
        'outros_acesso': '',
        'outros_ensino': '',
        'monitoramento_data': date.today(),
        'status_meta': 'Não Iniciado',
        'parecer_geral': 'Manter Estratégias',
        'proximos_passos_select': [],
        'status_validacao_pei': 'rascunho',
        'feedback_ajuste': '',
        'status_validacao_game': 'rascunho',
        'feedback_ajuste_game': ''
    }
    st.session_state.pdf_text = ""


# ==============================================================================
# 8. ESTILO VISUAL (mais seguro: não colore TODOS os botões)
# ==============================================================================
st.markdown("""
<style>
  @import url('https://fonts.googleapis.com/css2?family=Nunito:wght@400;600;700;800&display=swap');

  html, body, [class*="css"] {
    font-family: 'Nunito', sans-serif;
    color: #2D3748;
    background-color: #F7FAFC;
  }

  .block-container {
    padding-top: 1.5rem !important;
    padding-bottom: 5rem !important;
  }

  /* Tabs */
  div[data-baseweb="tab-border"], div[data-baseweb="tab-highlight"] { display:none !important; }

  .stTabs [data-baseweb="tab-list"] {
    gap: 8px;
    display: flex;
    flex-wrap: wrap !important;
    white-space: normal !important;
    overflow-x: visible !important;
    padding: 10px 5px;
    width: 100%;
  }

  .stTabs [data-baseweb="tab"] {
    height: 38px;
    border-radius: 20px !important;
    background-color: #FFFFFF;
    border: 1px solid #E2E8F0;
    color: #718096;
    font-weight: 700;
    font-size: 0.8rem;
    padding: 0 20px;
    box-shadow: 0 1px 2px rgba(0,0,0,0.03);
    text-transform: uppercase;
    letter-spacing: 0.5px;
    margin-bottom: 5px;
  }

  .stTabs [data-baseweb="tab"]:hover {
    border-color: #CBD5E0;
    color: #4A5568;
    background-color: #EDF2F7;
  }

  .stTabs [aria-selected="true"] {
    background-color: transparent !important;
    color: #3182CE !important;
    border: 1px solid #3182CE !important;
    font-weight: 800;
    box-shadow: 0 0 12px rgba(49,130,206,0.25), inset 0 0 5px rgba(49,130,206,0.08) !important;
  }

  /* Header */
  .header-unified {
    background-color: white;
    padding: 20px 28px;
    border-radius: 16px;
    border: 1px solid #E2E8F0;
    box-shadow: 0 2px 10px rgba(0,0,0,0.02);
    margin-bottom: 18px;
    display: flex;
    align-items: center;
    gap: 18px;
  }

  .header-subtitle {
    font-size: 1.12rem;
    color: #718096;
    font-weight: 700;
    border-left: 2px solid #E2E8F0;
    padding-left: 16px;
    line-height: 1.2;
  }

  /* Inputs */
  .stTextInput input, .stTextArea textarea,
  .stSelectbox div[data-baseweb="select"], .stMultiSelect div[data-baseweb="select"] {
    border-radius: 10px !important;
    border-color: #E2E8F0 !important;
  }

  /* Botões: NÃO aplicar globalmente.
     Só vamos “embelezar” os botões dentro de containers com classe .pei-actions */
  .pei-actions div[data-testid="stButton"] button {
    border-radius: 10px !important;
    font-weight: 800 !important;
    height: 44px !important;
    border: 0 !important;
  }

  .pei-actions .btn-primary div[data-testid="stButton"] button {
    background-color: #0F52BA !important;
    color: white !important;
  }
  .pei-actions .btn-primary div[data-testid="stButton"] button:hover {
    background-color: #0A3D8F !important;
  }

  .pei-actions .btn-ghost div[data-testid="stButton"] button {
    background-color: #EEF2FF !important;
    color: #1F2937 !important;
    border: 1px solid #E5E7EB !important;
  }
  .pei-actions .btn-ghost div[data-testid="stButton"] button:hover {
    background-color: #E5E7EB !important;
  }

  .footer-signature {
    text-align:center;
    opacity:0.55;
    font-size:0.75rem;
    padding:30px 0 10px 0;
  }
</style>

<link href="https://cdn.jsdelivr.net/npm/remixicon@4.1.0/fonts/remixicon.css" rel="stylesheet">
""", unsafe_allow_html=True)


# ==============================================================================
# 9. SIDEBAR — Sessão + OpenAI + Sincronização + Salvar/Carregar
#    (sem duplicar navegação / sem loops / com guard supabase)
# ==============================================================================
with st.sidebar:
    st.markdown("### 👤 Sessão")
    st.caption(f"Usuário: **{st.session_state.get('usuario_nome','')}**")
    st.caption(f"Workspace: **{st.session_state.get('workspace_name','')}**")

    st.markdown("---")
    st.markdown("### 🧾 Status do Aluno")
    student_id = st.session_state.get("selected_student_id")

    if student_id:
        st.success("✅ Vinculado ao Supabase")
        st.caption(f"student_id: {student_id[:8]}...")
    else:
        st.warning("📝 Rascunho (ainda não salvo na nuvem)")

    st.markdown("---")

    # OpenAI (prioriza secrets)
    if "OPENAI_API_KEY" in st.secrets and st.secrets.get("OPENAI_API_KEY"):
        api_key = st.secrets["OPENAI_API_KEY"]
        st.success("✅ OpenAI OK")
    else:
        api_key = st.text_input("Chave OpenAI:", type="password", key="pei_openai_key_input")

    st.markdown("---")
    st.markdown("### 🔗 Sincronização (Criar aluno)")

    # Se você ainda não tiver login supabase real, este bloco não quebra:
    has_sb_login = bool(st.session_state.get("supabase_jwt")) and bool(st.session_state.get("supabase_user_id"))
    if not has_sb_login:
        st.info("Faça login Supabase na Home para habilitar sincronização/salvar.")
    else:
        if not student_id:
            st.markdown('<div class="pei-actions btn-primary">', unsafe_allow_html=True)
            if st.button("🔗 Sincronizar agora (criar aluno)", use_container_width=True, key="sb_sync_create_student"):
                # validações mínimas
                if not st.session_state.dados.get("nome"):
                    st.warning("Preencha o NOME do estudante na aba Estudante antes de sincronizar.")
                elif not st.session_state.dados.get("serie"):
                    st.warning("Selecione a SÉRIE/Ano na aba Estudante antes de sincronizar.")
                else:
                    try:
                        created = db_create_student({
                            "name": st.session_state.dados.get("nome"),
                            "birth_date": st.session_state.dados.get("nasc").isoformat() if hasattr(st.session_state.dados.get("nasc"), "isoformat") else None,
                            "grade": st.session_state.dados.get("serie"),
                            "class_group": st.session_state.dados.get("turma") or None,
                            "diagnosis": st.session_state.dados.get("diagnostico") or None,
                        })
                        if created and created.get("id"):
                            st.session_state["selected_student_id"] = created["id"]
                            st.session_state["selected_student_name"] = created.get("name") or ""
                            st.success("Sincronizado ✅ Agora você pode salvar/carregar.")
                            st.rerun()
                        else:
                            st.error("Falha ao criar aluno. Verifique policies/RLS no Supabase.")
                    except Exception as e:
                        st.error(f"Erro ao sincronizar: {e}")
            st.markdown("</div>", unsafe_allow_html=True)
        else:
            st.caption("Aluno já sincronizado. Use Salvar/Carregar abaixo.")

    st.markdown("---")
    st.markdown("### 💾 Supabase (PEI)")

    c1, c2 = st.columns(2)

    with c1:
        st.markdown('<div class="pei-actions btn-primary">', unsafe_allow_html=True)
        if st.button("💾 Salvar", use_container_width=True, disabled=(not student_id), key="sb_save_pei_btn"):
            if not _sb_ok():
                st.error("Supabase não pronto (sb/owner/workspace).")
            else:
                with st.spinner("Salvando..."):
                    supa_save_pei(student_id, st.session_state.dados, st.session_state.get("pdf_text", ""))
                    supa_sync_student_from_dados(student_id, st.session_state.dados)
                st.success("Salvo no Supabase ✅")
        st.markdown("</div>", unsafe_allow_html=True)

    with c2:
        st.markdown('<div class="pei-actions btn-ghost">', unsafe_allow_html=True)
        if st.button("🔄 Carregar", use_container_width=True, disabled=(not student_id), key="sb_load_pei_btn"):
            if not _sb_ok():
                st.error("Supabase não pronto (sb/owner/workspace).")
            else:
                with st.spinner("Carregando..."):
                    row = supa_load_latest_pei(student_id)
                if row and row.get("payload"):
                    payload = row["payload"] or {}
                    # re-hidratar datas
                    for k in ["nasc", "monitoramento_data"]:
                        try:
                            if k in payload and isinstance(payload[k], str) and payload[k]:
                                payload[k] = date.fromisoformat(payload[k])
                        except:
                            pass
                    st.session_state.dados.update(payload)
                    st.session_state.pdf_text = row.get("pdf_text") or ""
                    st.success("Carregado ✅")
                    st.rerun()
                else:
                    st.info("Ainda não existe PEI salvo para este aluno.")
        st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("### 🧭 Navegação")
    # IMPORTANTE: aqui mantemos apenas Home (evita loops)
    if st.button("🏠 Home", use_container_width=True, key="nav_home_from_pei"):
        st.switch_page("streamlit_app.py")


# ==============================================================================
# 10. HEADER + ABAS (corrigido: 10 abas => tab0..tab9)
# ==============================================================================
logo_path = finding_logo()
b64_logo = get_base64_image(logo_path)
mime = "image/png"
img_html = f'<img src="data:{mime};base64,{b64_logo}" style="height: 96px;">' if logo_path else ""

st.markdown(
    f"""<div class="header-unified">
          {img_html}
          <div class="header-subtitle">Planejamento Educacional Inclusivo Inteligente</div>
        </div>""",
    unsafe_allow_html=True
)

abas = [
    "INÍCIO", "ESTUDANTE", "EVIDÊNCIAS", "REDE DE APOIO", "MAPEAMENTO",
    "PLANO DE AÇÃO", "MONITORAMENTO", "CONSULTORIA IA", "DASHBOARD & DOCS", "JORNADA GAMIFICADA"
]
tab0, tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8, tab9 = st.tabs(abas)

# ==============================================================================
# 11. ABA INÍCIO — CENTRAL (Gestão de Alunos + Backups)
# ==============================================================================
# ✅ Helpers locais (somente UI)
def _coerce_dates_in_payload(d: dict):
    """Converte campos de data salvos como string de volta para date."""
    if not isinstance(d, dict):
        return d
    for k in ["nasc", "monitoramento_data"]:
        try:
            if k in d and isinstance(d[k], str) and d[k]:
                d[k] = date.fromisoformat(d[k])
        except:
            pass
    return d

def _sb_ready_for_students():
    """Supabase pronto + owner + workspace (necessário para listar/excluir/criar)."""
    return _sb_ok()

def db_delete_student(student_id: str):
    """
    Exclui aluno do workspace atual.
    Se não houver FK cascade, removemos PEIs primeiro (best effort).
    """
    if not _sb_ready_for_students():
        raise RuntimeError("Supabase não pronto (sb/owner/workspace).")

    # 1) tenta apagar PEIs do aluno (best effort)
    try:
        sb().table("pei_documents") \
            .delete() \
            .eq("student_id", student_id) \
            .eq("workspace_id", WORKSPACE_ID) \
            .execute()
    except Exception:
        pass

    # 2) apaga aluno (scoped no workspace)
    sb().table("students") \
        .delete() \
        .eq("id", student_id) \
        .eq("workspace_id", WORKSPACE_ID) \
        .execute()


with tab0:
    st.markdown("### 🏛️ Central de Fundamentos e Gestão")
    st.caption("Aqui você gerencia alunos (nuvem/backup) e acessa os fundamentos do PEI.")

    # -------------------------
    # LAYOUT 2 COLUNAS
    # -------------------------
    col_left, col_right = st.columns([1.15, 0.85])

    # =========================
    # ESQUERDA: Fundamentos
    # =========================
    with col_left:
        with st.container(border=True):
            st.markdown("#### 📚 Fundamentos do PEI")
            st.markdown(
                """
- O **PEI** organiza o planejamento individualizado com foco em **barreiras** e **apoios**.
- A lógica é **equidade**: ajustar **acesso, ensino e avaliação**, sem baixar expectativas.
- Base: **LBI (Lei 13.146/2015)**, LDB e diretrizes de Educação Especial na Perspectiva Inclusiva.
                """
            )

        with st.container(border=True):
            st.markdown("#### 🧭 Como usar a Omnisfera")
            st.markdown(
                """
1) **Estudante**: identificação + contexto + laudo (opcional)  
2) **Evidências**: o que foi observado e como aparece na rotina  
3) **Mapeamento**: barreiras + nível de apoio + potências  
4) **Plano de Ação**: acesso/ensino/avaliação  
5) **Consultoria IA**: gerar o documento técnico (validação do educador)  
6) **Dashboard**: KPIs + exportações + sincronização  
                """
            )

    # =========================
    # DIREITA: Gestão de alunos
    # =========================
    with col_right:
        st.markdown("#### 👤 Gestão de Alunos")

        # Status vínculo
        student_id = st.session_state.get("selected_student_id")
        if student_id:
            st.success("✅ Aluno vinculado ao Supabase")
            st.caption(f"student_id: {student_id[:8]}...")
        else:
            st.warning("📝 Modo rascunho (não salvo na nuvem)")

        # ---------
        # (1) Carregar JSON do PC
        # ---------
        with st.container(border=True):
            st.markdown("##### 1) Carregar Backup (.JSON)")
            up_json = st.file_uploader(
                "Envie um arquivo .json",
                type="json",
                key="inicio_uploader_json",
            )
            if up_json:
                try:
                    djson = json.load(up_json)
                    djson = _coerce_dates_in_payload(djson)

                    if "dados" in st.session_state and isinstance(st.session_state.dados, dict):
                        st.session_state.dados.update(djson)
                    else:
                        st.session_state.dados = djson

                    st.success("Backup carregado ✅")
                    st.toast("Dados aplicados ao formulário.", icon="✅")
                    st.rerun()
                except Exception as e:
                    st.error(f"Erro ao ler JSON: {e}")

            st.caption("Dica: use isso para migrar dados entre máquinas ou restaurar um PEI salvo.")

        # ---------
        # (2) Nuvem: listar / carregar / excluir
        # ---------
        with st.container(border=True):
            st.markdown("##### 2) Nuvem (Supabase) — Alunos")

            if not _sb_ready_for_students():
                st.info("Supabase não pronto (sb/owner/workspace). Faça login na Home e valide o PIN.")
            else:
                # Buscar alunos (scoped por workspace)
                alunos = db_list_students()

                if not alunos:
                    st.info("Nenhum aluno salvo ainda.")
                else:
                    opcoes = []
                    mapa = {}

                    for a in alunos:
                        nome = a.get("name") or "Sem nome"
                        serie = a.get("grade") or "-"
                        turma = a.get("class_group") or "-"
                        label = f"{nome} — {serie} / {turma}"
                        opcoes.append(label)
                        mapa[label] = a

                    sel = st.selectbox(
                        "Selecione um aluno",
                        options=opcoes,
                        index=None,
                        placeholder="Escolha para carregar/excluir…",
                        key="inicio_select_aluno",
                    )

                    cA, cB, cC = st.columns(3)

                    # Carregar
                    with cA:
                        if st.button("☁️ Carregar", use_container_width=True, disabled=(not sel), key="inicio_btn_carregar"):
                            a = mapa.get(sel)
                            if not a:
                                st.error("Aluno inválido.")
                            else:
                                sid = a["id"]
                                # vincula
                                st.session_state["selected_student_id"] = sid
                                st.session_state["selected_student_name"] = a.get("name") or ""

                                # tenta carregar o último PEI salvo
                                row = supa_load_latest_pei(sid)
                                if row and row.get("payload"):
                                    payload = _coerce_dates_in_payload(row["payload"] or {})
                                    st.session_state.dados.update(payload)
                                    st.session_state.pdf_text = row.get("pdf_text") or ""
                                    st.success("Aluno + último PEI carregados ✅")
                                else:
                                    # se não houver PEI, preenche só cadastro básico
                                    st.session_state.dados["nome"] = a.get("name") or ""
                                    try:
                                        bd = a.get("birth_date")
                                        if isinstance(bd, str) and bd:
                                            st.session_state.dados["nasc"] = date.fromisoformat(bd)
                                    except:
                                        pass
                                    st.session_state.dados["serie"] = a.get("grade")
                                    st.session_state.dados["turma"] = a.get("class_group") or ""
                                    st.session_state.dados["diagnostico"] = a.get("diagnosis") or ""
                                    st.info("Aluno carregado. Ainda não existe PEI salvo para este aluno.")

                                st.rerun()

                    # Excluir (com confirmação)
                    with cB:
                        confirmar = st.checkbox("Confirmar exclusão", value=False, key="inicio_ck_confirm_del")
                        if st.button(
                            "🗑️ Excluir",
                            use_container_width=True,
                            disabled=(not sel) or (not confirmar),
                            type="primary",
                            key="inicio_btn_excluir",
                        ):
                            a = mapa.get(sel)
                            if not a:
                                st.error("Aluno inválido.")
                            else:
                                try:
                                    db_delete_student(a["id"])

                                    if st.session_state.get("selected_student_id") == a["id"]:
                                        st.session_state["selected_student_id"] = None
                                        st.session_state["selected_student_name"] = ""

                                    st.success("Aluno excluído ✅")
                                    st.rerun()
                                except Exception as e:
                                    st.error(f"Erro ao excluir: {e}")

                    # Novo rascunho
                    with cC:
                        if st.button("🧹 Novo (Rascunho)", use_container_width=True, key="inicio_btn_novo"):
                            st.session_state["selected_student_id"] = None
                            st.session_state["selected_student_name"] = ""

                            # volta pro default_state (sem perder as chaves)
                            st.session_state.dados = default_state.copy()
                            st.session_state.pdf_text = ""

                            st.toast("Voltando para rascunho…", icon="📝")
                            st.rerun()

        # ---------
        # Sincronizar aluno (criar na nuvem) — reforço da sidebar
        # ---------
        with st.container(border=True):
            st.markdown("##### 🔗 Sincronizar aluno (criar e vincular)")
            st.caption("Cria o aluno na tabela **students** e libera Salvar/Carregar do PEI.")

            if not _sb_ready_for_students():
                st.info("Supabase não pronto. Faça login na Home e valide o PIN.")
            else:
                if st.session_state.get("selected_student_id"):
                    st.success("Este aluno já está sincronizado ✅")
                else:
                    if st.button("🔗 Sincronizar agora", type="primary", use_container_width=True, key="inicio_btn_sync"):
                        if not st.session_state.dados.get("nome"):
                            st.warning("Preencha o NOME do estudante na aba Estudante antes de sincronizar.")
                        elif not st.session_state.dados.get("serie"):
                            st.warning("Selecione a SÉRIE/Ano na aba Estudante antes de sincronizar.")
                        else:
                            try:
                                created = db_create_student({
                                    "name": st.session_state.dados.get("nome"),
                                    "birth_date": st.session_state.dados.get("nasc").isoformat() if hasattr(st.session_state.dados.get("nasc"), "isoformat") else None,
                                    "grade": st.session_state.dados.get("serie"),
                                    "class_group": st.session_state.dados.get("turma") or None,
                                    "diagnosis": st.session_state.dados.get("diagnostico") or None,
                                })
                                if created and created.get("id"):
                                    st.session_state["selected_student_id"] = created["id"]
                                    st.session_state["selected_student_name"] = created.get("name") or ""
                                    st.success("Sincronizado ✅ Agora você pode salvar/carregar PEI.")
                                    st.rerun()
                                else:
                                    st.error("Falha ao criar aluno. Verifique policies/RLS no Supabase.")
                            except Exception as e:
                                st.error(f"Erro ao sincronizar: {e}")


# ==============================================================================
# 12. ABA ESTUDANTE
# ==============================================================================
# ✅ (A partir daqui, mantenha o seu conteúdo EXISTENTE das abas tab1..tab8,
#     porque você já enviou e está bem completo.)
#     Eu NÃO vou duplicar aqui para não gerar widgets repetidos.
#     O que segue abaixo é apenas a TROCA da Jornada para tab9.
# ==============================================================================


# ==============================================================================
# 20. ABA — JORNADA GAMIFICADA (AGORA EM tab9)
# ==============================================================================
with tab9:
    render_progresso()

    nome_aluno = st.session_state.dados.get("nome") or "Estudante"
    serie = st.session_state.dados.get("serie") or ""
    hiperfoco = st.session_state.dados.get("hiperfoco") or ""
    potencias = st.session_state.dados.get("potencias") or []
    pei_ok = bool(st.session_state.dados.get("ia_sugestao"))

    # Header visual
    seg_nome, seg_cor, seg_desc = ("Selecione a Série", "#CBD5E0", "Defina a série na aba Estudante.")
    try:
        if serie:
            seg_nome, seg_cor, seg_desc = get_segmento_info_visual(serie)
    except Exception:
        # fallback (se só existir get_segmento_info_visual_v2)
        try:
            seg_nome, seg_cor, seg_desc = get_segmento_info_visual_v2(serie)
        except Exception:
            pass

    st.markdown(
        f"""
        <div style="
            background: linear-gradient(90deg, {seg_cor} 0%, #111827 140%);
            padding: 22px 26px; border-radius: 18px; color: white; margin-bottom: 18px;
            box-shadow: 0 8px 18px rgba(0,0,0,0.06);
        ">
            <div style="display:flex; align-items:center; justify-content:space-between; gap:16px;">
                <div>
                    <div style="font-size:0.9rem; opacity:0.9; font-weight:800; letter-spacing:0.3px;">🎮 JORNADA GAMIFICADA</div>
                    <div style="font-size:1.6rem; font-weight:900; margin-top:4px;">Missão do(a) {nome_aluno}</div>
                    <div style="opacity:0.92; margin-top:6px; font-weight:700;">{seg_nome} • {serie}</div>
                </div>
                <div style="text-align:right;">
                    <div style="font-size:0.75rem; opacity:0.85; font-weight:800;">Modo</div>
                    <div style="font-size:1.05rem; font-weight:900;">{("Pronto" if pei_ok else "Aguardando PEI")}</div>
                </div>
            </div>
        </div>
        """,
        unsafe_allow_html=True
    )

    st.info(
        "ℹ️ Esta aba cria um material **para o estudante**: uma versão gamificada do plano, "
        "para imprimir, entregar à família ou usar como combinado de sala."
    )

    # Pré-requisitos
    if not serie:
        st.warning("⚠️ Selecione a **Série/Ano** na aba **Estudante** para liberar a Jornada.")
        st.stop()

    if not st.session_state.dados.get("nome"):
        st.warning("⚠️ Preencha o **nome do estudante** na aba **Estudante** para liberar a Jornada.")
        st.stop()

    if not pei_ok:
        st.warning("⚠️ Gere o PEI Técnico na aba **Consultoria IA** antes de criar a Jornada.")
        st.stop()

    # Contexto compacto
    with st.container(border=True):
        cA, cB, cC = st.columns([2, 2, 2])
        with cA:
            st.markdown("##### 🚀 Hiperfoco")
            st.write(hiperfoco if hiperfoco else "—")
        with cB:
            st.markdown("##### 🌟 Potencialidades")
            st.write(", ".join(potencias) if potencias else "—")
        with cC:
            st.markdown("##### 🧭 Guia do Segmento")
            st.caption(seg_desc)

    st.divider()

    # Estado de validação
    st.session_state.dados.setdefault("status_validacao_game", "rascunho")
    st.session_state.dados.setdefault("feedback_ajuste_game", "")
    st.session_state.dados.setdefault("ia_mapa_texto", "")

    status_game = st.session_state.dados.get("status_validacao_game", "rascunho")

    st.markdown("### 🧩 Gerar / Revisar Missão")

    # 1) RASCUNHO — gerar
    if status_game == "rascunho":
        st.markdown(
            """
**Como funciona**
- A IA usa **hiperfoco + potências** para criar uma história motivadora.
- O texto evita dados sensíveis e foca em **apoio, autonomia e rotina**.
            """
        )

        col1, col2 = st.columns([2, 1])
        with col1:
            estilo = st.text_input(
                "Preferência de estilo (opcional)",
                placeholder="Ex: super-heróis, exploração espacial, futebol, fantasia medieval...",
                key="gm_estilo_tab9",
            )
        with col2:
            st.write("")
            gerar_btn = st.button("🎮 Criar Roteiro Gamificado", type="primary", use_container_width=True, key="gm_btn_gerar_tab9")

        if gerar_btn:
            with st.spinner("Game Master criando a missão..."):
                fb = (f"Estilo desejado: {estilo}." if estilo else "").strip()
                texto_game, err = gerar_roteiro_gamificado(
                    api_key,
                    st.session_state.dados,
                    st.session_state.dados.get("ia_sugestao", ""),
                    fb
                )

                if texto_game:
                    st.session_state.dados["ia_mapa_texto"] = texto_game.replace("[MAPA_TEXTO_GAMIFICADO]", "").strip()
                    st.session_state.dados["status_validacao_game"] = "revisao"
                    st.rerun()
                else:
                    st.error(err or "Erro desconhecido ao gerar a missão.")

    # 2) REVISÃO — aprovar/refazer
    elif status_game == "revisao":
        st.success("✅ Missão gerada! Revise abaixo e aprove/solicite ajustes.")

        with st.container(border=True):
            st.markdown("#### 📜 Missão (prévia)")
            st.markdown(st.session_state.dados.get("ia_mapa_texto", ""))

        st.divider()
        c_ok, c_aj = st.columns(2)
        with c_ok:
            if st.button("✅ Aprovar Missão", type="primary", use_container_width=True, key="gm_btn_aprovar_tab9"):
                st.session_state.dados["status_validacao_game"] = "aprovado"
                st.rerun()
        with c_aj:
            if st.button("✏️ Solicitar Ajustes", use_container_width=True, key="gm_btn_ajustes_tab9"):
                st.session_state.dados["status_validacao_game"] = "ajustando"
                st.rerun()

    # 3) AJUSTANDO — feedback e regerar
    elif status_game == "ajustando":
        st.warning("🛠️ Descreva o que você quer mudar e regenere a missão.")

        fb_game = st.text_area(
            "O que ajustar na missão?",
            value=st.session_state.dados.get("feedback_ajuste_game", ""),
            placeholder="Ex: deixe mais curto, use linguagem mais infantil, traga recompensas, troque o tema para futebol...",
            height=140,
            key="gm_txt_feedback_tab9"
        )
        st.session_state.dados["feedback_ajuste_game"] = fb_game

        c1, c2 = st.columns([2, 1])
        with c1:
            if st.button("🔁 Regerar com Ajustes", type="primary", use_container_width=True, key="gm_btn_regerar_tab9"):
                with st.spinner("Reescrevendo missão..."):
                    texto_game, err = gerar_roteiro_gamificado(
                        api_key,
                        st.session_state.dados,
                        st.session_state.dados.get("ia_sugestao", ""),
                        feedback_game=fb_game
                    )
                    if texto_game:
                        st.session_state.dados["ia_mapa_texto"] = texto_game.replace("[MAPA_TEXTO_GAMIFICADO]", "").strip()
                        st.session_state.dados["status_validacao_game"] = "revisao"
                        st.rerun()
                    else:
                        st.error(err or "Erro desconhecido ao regerar a missão.")
        with c2:
            if st.button("↩️ Voltar", use_container_width=True, key="gm_btn_voltar_tab9"):
                st.session_state.dados["status_validacao_game"] = "revisao"
                st.rerun()

    # 4) APROVADO — exportar PDF e editar fino
    elif status_game == "aprovado":
        st.success("🏁 Missão aprovada! Agora você pode imprimir e entregar.")

        colA, colB = st.columns([2, 1])
        with colA:
            with st.container(border=True):
                st.markdown("#### 📜 Missão Final (editável)")
                novo_texto = st.text_area(
                    "Edição final manual (opcional)",
                    value=st.session_state.dados.get("ia_mapa_texto", ""),
                    height=320,
                    key="gm_txt_final_tab9"
                )
                st.session_state.dados["ia_mapa_texto"] = novo_texto

        with colB:
            with st.container(border=True):
                st.markdown("#### 📥 Exportação")
                pdf_mapa = gerar_pdf_tabuleiro_simples(st.session_state.dados["ia_mapa_texto"])
                st.download_button(
                    "📄 Baixar Missão em PDF",
                    pdf_mapa,
                    file_name=f"Missao_{nome_aluno}.pdf",
                    mime="application/pdf",
                    type="primary",
                    use_container_width=True,
                    key="gm_dl_pdf_tab9"
                )
                st.caption("Dica: imprima e cole no caderno / agenda do aluno.")
                st.write("---")
                if st.button("🆕 Criar Nova Missão", use_container_width=True, key="gm_btn_nova_tab9"):
                    st.session_state.dados["status_validacao_game"] = "rascunho"
                    st.session_state.dados["feedback_ajuste_game"] = ""
                    st.session_state.dados["ia_mapa_texto"] = ""
                    st.rerun()

    else:
        st.session_state.dados["status_validacao_game"] = "rascunho"
        st.rerun()
