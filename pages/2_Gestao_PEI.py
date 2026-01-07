import streamlit as st
from datetime import date
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openai import OpenAI
from pypdf import PdfReader
from fpdf import FPDF
import base64
import os
import re

# --- CONFIGURAÇÃO DA PÁGINA ---
st.set_page_config(
    page_title="PEI 360º | Sistema Inclusivo",
    page_icon="📘",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- 0. INICIALIZAÇÃO DO BANCO DE DADOS ---
if 'banco_estudantes' not in st.session_state:
    st.session_state.banco_estudantes = []

# --- FUNÇÃO FAVICON ---
def get_favicon():
    if os.path.exists("iconeaba.png"): return "iconeaba.png"
    if os.path.exists("360.png"): return "360.png"
    return "📘"

# --- ESTILO VISUAL CORAL & BLUE (ARCO PALETTE) ---
st.markdown("""
    <link href="https://cdn.jsdelivr.net/npm/remixicon@4.1.0/fonts/remixicon.css" rel="stylesheet">
    <link href="https://fonts.googleapis.com/css2?family=Nunito:wght@400;600;700;800&display=swap" rel="stylesheet">
    
    <style>
    /* 1. GLOBAL & CORES */
    html, body, [class*="css"] { font-family: 'Nunito', sans-serif; color: #2D3748; }
    :root { 
        --brand-blue: #004E92;
        --brand-coral: #FF6B6B;
        --bg-light: #F7FAFC; 
        --card-shadow: 0 4px 6px rgba(0,0,0,0.04);
    }
    
    /* 2. HEADER */
    .header-container {
        padding: 25px; background: #FFFFFF; border-radius: 20px; 
        border: 1px solid #EDF2F7; border-left: 8px solid var(--brand-blue); 
        box-shadow: var(--card-shadow); margin-bottom: 30px;
        display: flex; align-items: center; gap: 25px;
    }
    
    /* 3. ABAS */
    .stTabs [data-baseweb="tab-list"] {
        gap: 10px; background-color: transparent; padding: 10px 0;
        justify-content: flex-start; flex-wrap: wrap;
    }
    .stTabs [data-baseweb="tab"] {
        height: 42px; background-color: #FFFFFF; border-radius: 20px;
        border: 1px solid #CBD5E0; color: #4A5568; padding: 0 20px;
        font-weight: 700; font-size: 0.9rem; flex-grow: 0; transition: all 0.2s ease;
    }
    .stTabs [aria-selected="true"] {
        background-color: var(--brand-coral) !important;
        color: white !important;
        border-color: var(--brand-coral) !important;
        box-shadow: 0 4px 10px rgba(255, 107, 107, 0.3);
    }

    /* 4. CARDS */
    .feature-card {
        background: white; padding: 25px; border-radius: 20px;
        border: 1px solid #EDF2F7; box-shadow: var(--card-shadow);
        height: 100%; transition: all 0.3s ease;
        display: flex; flex-direction: column; align-items: flex-start;
    }
    .icon-box {
        width: 45px; height: 45px; background: #E3F2FD; border-radius: 12px;
        display: flex; align-items: center; justify-content: center; margin-bottom: 15px; flex-shrink: 0;
    }
    .icon-box i { font-size: 22px; color: var(--brand-blue); }
    .feature-card h4 { color: var(--brand-blue); font-weight: 800; font-size: 1.1rem; margin-bottom: 8px; line-height: 1.3; }
    .feature-card p { font-size: 0.95rem; color: #718096; line-height: 1.5; margin: 0; }

    /* 5. INPUTS & BOTÕES */
    .stTextInput input, .stTextArea textarea, .stSelectbox div[data-baseweb="select"] {
        border-radius: 12px !important; border: 1px solid #CBD5E0 !important;
    }
    .stTextInput input:focus, .stTextArea textarea:focus {
        border-color: var(--brand-blue) !important; box-shadow: 0 0 0 2px rgba(0, 78, 146, 0.2) !important;
    }
    
    div[data-testid="column"] .stButton button[kind="primary"] {
        background-color: var(--brand-coral) !important; 
        color: white !important; border: none !important; border-radius: 12px !important; 
        font-weight: 700 !important; height: 3.5em !important; width: 100%; transition: 0.3s !important;
    }
    div[data-testid="column"] .stButton button[kind="primary"]:hover {
        background-color: #E53E3E !important; transform: scale(1.02) !important;
    }
    div[data-testid="column"] .stButton button[kind="secondary"] {
        background-color: transparent !important; color: var(--brand-blue) !important;
        border: 2px solid var(--brand-blue) !important; border-radius: 12px !important; 
        font-weight: 700 !important; height: 3.5em !important; width: 100%;
    }
    span[data-baseweb="tag"] { background-color: #EBF8FF !important; border: 1px solid #90CDF4 !important; }
    span[data-baseweb="tag"] span { color: #004E92 !important; }
    </style>
    """, unsafe_allow_html=True)

# --- FUNÇÕES AUXILIARES ---
def finding_logo():
    possiveis = ["360.png", "360.jpg", "logo.png", "logo.jpg"]
    for nome in possiveis:
        if os.path.exists(nome): return nome
    return None

def get_base64_image(image_path):
    if not image_path: return ""
    with open(image_path, "rb") as img_file:
        return base64.b64encode(img_file.read()).decode()

def ler_pdf(arquivo):
    if arquivo is None: return ""
    try:
        reader = PdfReader(arquivo)
        texto = ""
        for page in reader.pages: texto += page.extract_text() + "\n"
        return texto
    except Exception as e: return f"Erro: {e}"

def limpar_markdown(texto):
    if not texto: return ""
    texto = texto.replace('**', '').replace('__', '')
    texto = texto.replace('### ', '').replace('## ', '').replace('# ', '')
    return texto

def limpar_para_pdf(texto):
    if not texto: return ""
    texto = texto.replace('**', '').replace('__', '')
    texto = texto.replace('### ', '').replace('## ', '').replace('# ', '')
    texto = texto.replace('* ', '• ')
    texto = re.sub(r'[^\x00-\x7F\xA0-\xFF]', '', texto) 
    return texto

def calcular_idade(data_nasc):
    if not data_nasc: return ""
    hoje = date.today()
    idade = hoje.year - data_nasc.year - ((hoje.month, hoje.day) < (data_nasc.month, data_nasc.day))
    return str(idade)

# --- INTELIGÊNCIA REFINADA (COM GERAÇÃO DE PERFIL BNCC/IDADE) ---
def consultar_ia(api_key, dados, contexto_pdf=""):
    if not api_key: return None, "⚠️ A chave de API não foi detectada."
    try:
        client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com")
        serie = dados['serie'] if dados['serie'] else ""
        idade = calcular_idade(dados.get('nasc'))
        foco_bncc = "Campos de Experiência" if "Infantil" in serie else "Habilidades Essenciais"

        prompt_sistema = """
        Você é um Especialista Sênior em Inclusão, Neurociência e BNCC.
        Sua missão é gerar um PEI estratégico e um Dossiê Técnico para Adaptação de Provas.
        """
        
        contexto_extra = f"\n📄 LAUDO:{contexto_pdf[:3000]}" if contexto_pdf else ""
        
        prompt_usuario = f"""
        Estudante: {dados['nome']} | Idade Real: {idade} anos | Série: {serie}
        Diagnóstico: {dados['diagnostico']} | Hiperfoco: {dados['hiperfoco']}
        
        HISTÓRICO: {dados['historico']}
        FAMÍLIA: {dados['familia']}
        BARREIRAS: {', '.join(dados['b_sensorial'] + dados['b_cognitiva'] + dados['b_social'])}
        ESTRATÉGIAS: {', '.join(dados['estrategias_acesso'] + dados['estrategias_ensino'])}
        {contexto_extra}
        
        GERE O TEXTO NESTA ESTRUTURA ESTRATÉGICA:
        
        1. SÍNTESE DO CONTEXTO
        (Resumo conectando histórico, idade e diagnóstico).
        
        2. ANÁLISE NEUROFUNCIONAL
        (Como o cérebro aprende usando o Hiperfoco).
        
        3. ESTRATÉGIA BNCC ({foco_bncc})
        (Cite 1 objetivo da série e a flexibilização necessária).
        
        4. RECOMENDAÇÕES DE ROTINA
        (Dicas práticas de sala).

        5. DIRETRIZES PARA O ADAPTADOR DE PROVAS (Importante!)
        (Escreva um parágrafo técnico instruindo uma IA futura sobre como adaptar provas para este aluno específico. Cite o Nível de Leitura esperado para a idade/série, o tipo de fonte ideal e se deve usar imagens. Ex: "Para este aluno do 3º ano com TEA, as provas devem ter enunciados diretos, fonte Arial 14 e apoio visual para interpretação de texto.")
        """
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "system", "content": prompt_sistema}, {"role": "user", "content": prompt_usuario}],
            temperature=0.7, stream=False
        )
        return response.choices[0].message.content, None
    except Exception as e: return None, f"Erro DeepSeek: {str(e)}"

# --- PDF ---
class PDF(FPDF):
    def header(self):
        logo = finding_logo()
        if logo:
            self.image(logo, x=10, y=8, w=25)
            x = 40
        else: x = 10
        self.set_font('Arial', 'B', 16); self.set_text_color(0, 78, 146)
        self.cell(x); self.cell(0, 10, 'PLANO DE ENSINO INDIVIDUALIZADO', 0, 1, 'C'); self.ln(5)
    def footer(self):
        self.set_y(-15); self.set_font('Arial', 'I', 8); self.set_text_color(128)
        self.cell(0, 10, f'Página {self.page_no()} | Documento Confidencial', 0, 0, 'C')

def gerar_pdf_nativo(dados):
    pdf = PDF(); pdf.add_page(); pdf.set_font("Arial", size=11)
    def txt(t): return str(t).encode('latin-1', 'replace').decode('latin-1')

    # Identificação com Idade Calculada
    idade = calcular_idade(dados.get('nasc'))
    pdf.set_font("Arial", 'B', 12); pdf.set_text_color(0, 78, 146)
    pdf.cell(0, 10, txt("1. IDENTIFICAÇÃO DO ESTUDANTE"), 0, 1)
    pdf.set_font("Arial", size=11); pdf.set_text_color(0)
    
    nasc = dados.get('nasc'); d_nasc = nasc.strftime('%d/%m/%Y') if nasc else "-"
    pdf.multi_cell(0, 7, txt(f"Nome: {dados['nome']} | Idade: {idade} anos | Série: {dados['serie']}\nNascimento: {d_nasc}\nDiagnóstico: {dados['diagnostico']}"))
    pdf.ln(3)

    # Estratégias
    pdf.set_font("Arial", 'B', 12); pdf.set_text_color(0, 78, 146)
    pdf.cell(0, 10, txt("2. ESTRATÉGIAS EDUCACIONAIS"), 0, 1)
    pdf.set_font("Arial", size=11); pdf.set_text_color(0)
    
    if dados['estrategias_acesso']:
        pdf.set_font("Arial", 'B', 11); pdf.cell(0, 8, txt("Acesso e Organização:"), 0, 1); pdf.set_font("Arial", size=11)
        pdf.multi_cell(0, 7, txt(limpar_para_pdf(', '.join(dados['estrategias_acesso']))))
    
    if dados['estrategias_ensino']:
        pdf.set_font("Arial", 'B', 11); pdf.cell(0, 8, txt("Metodologia e Ensino:"), 0, 1); pdf.set_font("Arial", size=11)
        pdf.multi_cell(0, 7, txt(limpar_para_pdf(', '.join(dados['estrategias_ensino']))))
    
    # Parecer IA
    if dados['ia_sugestao']:
        pdf.ln(5)
        pdf.set_font("Arial", 'B', 12); pdf.set_text_color(0, 78, 146)
        pdf.cell(0, 10, txt("3. PARECER TÉCNICO PEDAGÓGICO"), 0, 1)
        pdf.set_font("Arial", size=11); pdf.set_text_color(50)
        conteudo_ia = limpar_para_pdf(dados['ia_sugestao'])
        pdf.multi_cell(0, 6, txt(conteudo_ia))

    pdf.ln(15); pdf.set_draw_color(0); pdf.line(20, pdf.get_y(), 190, pdf.get_y())
    pdf.cell(0, 10, txt("Coordenação Pedagógica / Direção Escolar"), 0, 1, 'C')
    return pdf.output(dest='S').encode('latin-1')

def gerar_docx_final(dados):
    doc = Document(); style = doc.styles['Normal']; style.font.name = 'Arial'; style.font.size = Pt(11)
    doc.add_heading('PLANO DE ENSINO INDIVIDUALIZADO', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Nome: {dados['nome']} | Série: {dados['serie']}")
    
    doc.add_heading('Estratégias', level=1)
    doc.add_paragraph(f"Acesso: {', '.join(dados['estrategias_acesso'])}")
    doc.add_paragraph(f"Ensino: {', '.join(dados['estrategias_ensino'])}")
    
    if dados['ia_sugestao']:
        doc.add_heading('Parecer Técnico', level=1)
        doc.add_paragraph(limpar_markdown(dados['ia_sugestao']))
    buffer = BytesIO(); doc.save(buffer); buffer.seek(0)
    return buffer

# --- ESTADO INICIAL ---
if 'dados' not in st.session_state:
    st.session_state.dados = {
        'nome': '', 'nasc': None, 'serie': None, 'escola': '', 'tem_laudo': False, 'diagnostico': '', 
        'rede_apoio': [], 'historico': '', 'familia': '', 'hiperfoco': '', 'potencias': [], 
        'b_sensorial': [], 'sup_sensorial': '🟡 Monitorado',
        'b_cognitiva': [], 'sup_cognitiva': '🟡 Monitorado',
        'b_social': [], 'sup_social': '🟡 Monitorado',
        'estrategias_acesso': [], 'estrategias_ensino': [], 'estrategias_avaliacao': [],
        'ia_sugestao': ''
    }
for k in ['estrategias_ensino', 'estrategias_avaliacao', 'rede_apoio']:
    if k not in st.session_state.dados: st.session_state.dados[k] = []
if 'nasc' not in st.session_state.dados: st.session_state.dados['nasc'] = None
if 'pdf_text' not in st.session_state: st.session_state.pdf_text = ""

# --- SIDEBAR ---
with st.sidebar:
    logo = finding_logo()
    if logo: st.image(logo, width=120)
    if 'DEEPSEEK_API_KEY' in st.secrets:
        api_key = st.secrets['DEEPSEEK_API_KEY']; st.success("✅ Chave Segura")
    else: api_key = st.text_input("Chave API:", type="password")
    st.markdown("---"); st.info("Versão 2.19 | Integration Ready")

# --- CABEÇALHO ---
logo = finding_logo()
header_html = ""
if logo:
    mime = "image/png" if logo.lower().endswith("png") else "image/jpeg"
    b64 = get_base64_image(logo)
    header_html = f"""
    <div class="header-container">
        <img src="data:{mime};base64,{b64}" class="header-logo" style="max-height: 105px; width: auto;"> 
        <div class="header-text" style="border-left: 2px solid #E2E8F0; padding-left: 25px;">
            <p style="margin: 0; color: #004E92; font-weight: 700; font-size: 1.2rem;">Planejamento Educacional Individualizado</p>
        </div>
    </div>
    """
else:
    header_html = '<div style="padding: 25px; background: white; border-radius: 20px; border: 1px solid #EDF2F7; box-shadow: 0 4px 6px rgba(0,0,0,0.02); margin-bottom: 30px;"><h1 style="color: #004E92; margin: 0;">PEI 360º</h1></div>'
st.markdown(header_html, unsafe_allow_html=True)

# ABAS
abas = ["Início", "Estudante", "Mapeamento", "Plano de Ação", "Assistente de IA", "Documento", "💾 Salvar no Sistema"]
tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs(abas)

# 1. HOME
with tab1:
    st.markdown("### <i class='ri-dashboard-line'></i> Ecossistema de Inclusão", unsafe_allow_html=True)
    c1, c2 = st.columns(2)
    with c1:
        st.markdown('<div class="feature-card"><div class="icon-box"><i class="ri-book-open-line"></i></div><h4>O que é o PEI?</h4><p>Acessibilidade oficial para flexibilizar o ensino.</p></div>', unsafe_allow_html=True)
    with c2:
        st.markdown('<div class="feature-card"><div class="icon-box"><i class="ri-scales-3-line"></i></div><h4>Legislação</h4><p>Documento obrigatório para garantir o suporte.</p></div>', unsafe_allow_html=True)

# 2. ESTUDANTE
with tab2:
    st.markdown("### <i class='ri-user-3-line'></i> Dossiê do Estudante", unsafe_allow_html=True)
    c1, c2, c3 = st.columns([2, 1, 1])
    st.session_state.dados['nome'] = c1.text_input("Nome do Estudante", st.session_state.dados['nome'])
    st.session_state.dados['nasc'] = c2.date_input("Data de Nascimento", st.session_state.dados.get('nasc'), format="DD/MM/YYYY")
    st.session_state.dados['serie'] = c3.selectbox("Série/Ano", ["Ed. Infantil", "1º Ano", "2º Ano", "3º Ano", "4º Ano", "5º Ano", "6º Ano", "7º Ano", "8º Ano", "9º Ano", "Ensino Médio"], index=None)
    
    st.markdown("---")
    ch, cf = st.columns(2)
    st.session_state.dados['historico'] = ch.text_area("Histórico Escolar", st.session_state.dados['historico'])
    st.session_state.dados['familia'] = cf.text_area("Escuta da Família", st.session_state.dados['familia'])
    
    st.markdown("---")
    c_diag, c_rede = st.columns(2)
    st.session_state.dados['diagnostico'] = c_diag.text_input("Diagnóstico", st.session_state.dados['diagnostico'])
    st.session_state.dados['rede_apoio'] = c_rede.multiselect("Rede de Apoio", ["Psicólogo", "Fonoaudiólogo", "Neuropediatra", "TO", "Psicopedagogo", "AT"], default=st.session_state.dados.get('rede_apoio', []))
    
    with st.expander("📂 Anexar Laudo (PDF)"):
        uploaded = st.file_uploader("Upload", type="pdf")
        if uploaded: st.session_state.pdf_text = ler_pdf(uploaded); st.success("Laudo lido!")

# 3. MAPEAMENTO
with tab3:
    st.markdown("### <i class='ri-rocket-line'></i> Potencialidades", unsafe_allow_html=True)
    st.session_state.dados['hiperfoco'] = st.text_input("Hiperfoco")
    
    with st.expander("Sensorial e Cognitivo", expanded=True):
        st.session_state.dados['b_sensorial'] = st.multiselect("Barreiras Sensoriais", ["Hipersensibilidade", "Busca Sensorial", "Seletividade"], key="b_sens")
        st.session_state.dados['b_cognitiva'] = st.multiselect("Barreiras Cognitivas", ["Atenção Dispersa", "Memória Curta", "Rigidez"], key="b_cog")

# 4. PLANO
with tab4:
    st.markdown("### <i class='ri-checkbox-circle-line'></i> Estratégias", unsafe_allow_html=True)
    c_a, c_b = st.columns(2)
    with c_a: st.session_state.dados['estrategias_acesso'] = st.multiselect("Acesso:", ["Tempo estendido", "Ledor", "Material Ampliado", "Sala Silenciosa"], key="acc")
    with c_b: st.session_state.dados['estrategias_ensino'] = st.multiselect("Ensino:", ["Pistas Visuais", "Mapa Mental", "Fragmentação", "Enunciados Curtos"], key="ens")

# 5. IA
with tab5:
    st.markdown("### <i class='ri-robot-line'></i> Consultor Inteligente", unsafe_allow_html=True)
    if st.button("✨ Gerar Parecer Completo (+ Diretrizes de Adaptação)", type="primary"):
        if not st.session_state.dados['nome']: st.warning("Preencha o nome.")
        else:
            with st.spinner("Gerando Dossiê Completo (Histórico + BNCC + Adaptação)..."):
                res, err = consultar_ia(api_key, st.session_state.dados, st.session_state.pdf_text)
                if err: st.error(err)
                else: st.session_state.dados['ia_sugestao'] = res; st.success("Parecer Gerado!")
    
    if st.session_state.dados['ia_sugestao']:
        st.markdown(f"<div style='background:white; padding:20px; border-radius:10px; border:1px solid #E2E8F0;'>{st.session_state.dados['ia_sugestao'].replace(chr(10), '<br>')}</div>", unsafe_allow_html=True)

# 6. DOCS
with tab6:
    st.markdown("<div style='text-align:center; padding: 30px;'>", unsafe_allow_html=True)
    if st.session_state.dados['nome']:
        c1, c2 = st.columns(2)
        with c1: st.download_button("📥 Baixar Word", gerar_docx_final(st.session_state.dados), "pei.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="secondary")
        with c2: st.download_button("📄 Baixar PDF", gerar_pdf_nativo(st.session_state.dados), "pei.pdf", "application/pdf", type="primary")

# 7. SALVAR (INTEGRAÇÃO INTELIGENTE)
with tab7:
    st.markdown("### <i class='ri-save-3-line'></i> Integração com o Adaptador de Provas", unsafe_allow_html=True)
    
    idade_calc = calcular_idade(st.session_state.dados.get('nasc'))
    serie_aluno = st.session_state.dados.get('serie', 'Não Informada')
    
    st.info(f"Ao salvar, criaremos um **Dossiê Técnico** para o aluno **{st.session_state.dados['nome']}** ({idade_calc} anos, {serie_aluno}), contendo as Diretrizes de Adaptação geradas pela IA.")
    
    if not st.session_state.dados.get('ia_sugestao'):
        st.warning("⚠️ Atenção: O Parecer da IA ainda não foi gerado na Aba 5. Recomenda-se gerar antes de salvar para ter as diretrizes completas.")
    
    if st.button("💾 Salvar Aluno e Enviar para o Adaptador", type="primary"):
        if st.session_state.dados['nome']:
            # Cria o objeto rico para o Adaptador
            perfil_rico = st.session_state.dados.copy()
            perfil_rico['idade_calculada'] = idade_calc
            # Salva no banco de sessão
            st.session_state.banco_estudantes.append(perfil_rico)
            st.success(f"Sucesso! O Dossiê de **{perfil_rico['nome']}** foi enviado para o Adaptador de Provas com todas as diretrizes da BNCC.")
        else:
            st.warning("Preencha pelo menos o nome.")

# --- RODAPÉ ---
st.markdown("""
<div style="text-align: center; margin-top: 50px; color: #A0AEC0; font-size: 0.85rem; border-top: 1px solid #E2E8F0; padding-top: 20px;">
    Criado e desenvolvido por Rodrigo Queiroz | Versão 2.19
</div>
""", unsafe_allow_html=True)
