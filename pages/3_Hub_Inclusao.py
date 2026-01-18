import streamlit as st
import os
from openai import OpenAI
from datetime import datetime, date
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from pypdf import PdfReader
from fpdf import FPDF
import base64
import re
import json
import requests
from PIL import Image
from streamlit_cropper import st_cropper
import gspread
from google.oauth2.service_account import Credentials

# ==============================================================================
# 1. CONFIGURAÇÃO E SEGURANÇA
# ==============================================================================
st.set_page_config(page_title="Hub de Inclusão | Omnisfera", page_icon="🚀", layout="wide")

# ==============================================================================
# 2. BLOCO VISUAL E UTILITÁRIOS
# ==============================================================================
try:
    IS_TEST_ENV = st.secrets.get("ENV") == "TESTE"
except:
    IS_TEST_ENV = False

def get_logo_base64():
    caminhos = ["omni_icone.png", "logo.png", "iconeaba.png"]
    for c in caminhos:
        if os.path.exists(c):
            with open(c, "rb") as f:
                return f"data:image/png;base64,{base64.b64encode(f.read()).decode()}"
    return "https://cdn-icons-png.flaticon.com/512/1183/1183672.png"

src_logo_giratoria = get_logo_base64()

if IS_TEST_ENV:
    card_bg, card_border = "rgba(255, 220, 50, 0.95)", "rgba(200, 160, 0, 0.5)"
else:
    card_bg, card_border = "rgba(255, 255, 255, 0.85)", "rgba(255, 255, 255, 0.6)"

st.markdown(f"""
<style>
    .omni-badge {{
        position: fixed; top: 15px; right: 15px;
        background: {card_bg}; border: 1px solid {card_border};
        backdrop-filter: blur(8px); padding: 4px 30px; min-width: 260px;
        justify-content: center; border-radius: 20px;
        box-shadow: 0 4px 15px rgba(0,0,0,0.08); z-index: 999990;
        display: flex; align-items: center; gap: 10px; pointer-events: none;
    }}
    .omni-text {{ font-family: 'Nunito', sans-serif; font-weight: 800; font-size: 0.9rem; color: #2D3748; letter-spacing: 1px; text-transform: uppercase; }}
    @keyframes spin-slow {{ from {{ transform: rotate(0deg); }} to {{ transform: rotate(360deg); }} }}
    .omni-logo-spin {{ height: 26px; width: 26px; animation: spin-slow 10s linear infinite; }}
    
    /* CSS Geral do Hub */
    html, body, [class*="css"] {{ font-family: 'Nunito', sans-serif; color: #2D3748; }}
    .header-hub {{ background: white; padding: 20px 30px; border-radius: 12px; border-left: 6px solid #3182CE; box-shadow: 0 2px 4px rgba(0,0,0,0.05); margin-bottom: 20px; display: flex; align-items: center; gap: 25px; }}
    .student-header {{ background-color: #EBF8FF; border: 1px solid #BEE3F8; border-radius: 12px; padding: 15px 25px; margin-bottom: 25px; display: flex; justify-content: space-between; align-items: center; }}
    .student-label {{ font-size: 0.85rem; color: #718096; font-weight: 700; text-transform: uppercase; }}
    .student-value {{ font-size: 1.1rem; color: #2C5282; font-weight: 800; }}
    .analise-box {{ background-color: #F0FFF4; border: 1px solid #C6F6D5; border-radius: 8px; padding: 20px; margin-bottom: 20px; color: #22543D; }}
    .validado-box {{ background-color: #C6F6D5; color: #22543D; padding: 15px; border-radius: 8px; text-align: center; font-weight: bold; margin-top: 15px; border: 1px solid #276749; }}
    .pedagogia-box {{ background-color: #F7FAFC; border-left: 4px solid #3182CE; padding: 15px; border-radius: 0 8px 8px 0; margin-bottom: 20px; font-size: 0.9rem; color: #4A5568; }}
    .pedagogia-title {{ color: #2C5282; font-weight: 700; display: flex; align-items: center; gap: 8px; margin-bottom: 5px; }}
    .stTabs [data-baseweb="tab-list"] {{ gap: 8px; flex-wrap: wrap; }}
    .stTabs [data-baseweb="tab"] {{ border-radius: 6px; padding: 8px 16px; background-color: white; border: 1px solid #E2E8F0; font-size: 0.9rem; transition: all 0.2s; }}
    .stTabs [aria-selected="true"] {{ background-color: #3182CE !important; color: white !important; border-color: #3182CE !important; }}
    div[data-testid="column"] .stButton button[kind="primary"] {{ border-radius: 10px !important; height: 50px; width: 100%; background-color: #3182CE !important; color: white !important; font-weight: 800 !important; border: none; }}
    div[data-testid="column"] .stButton button[kind="secondary"] {{ border-radius: 10px !important; height: 50px; width: 100%; border: 2px solid #CBD5E0 !important; color: #4A5568 !important; font-weight: bold; }}
</style>
<div class="omni-badge"><img src="{src_logo_giratoria}" class="omni-logo-spin"><span class="omni-text">OMNISFERA</span></div>
""", unsafe_allow_html=True)

def verificar_acesso():
    if "autenticado" not in st.session_state or not st.session_state["autenticado"]:
        pass 
    st.markdown("""<style>footer {visibility: hidden !important;} [data-testid="stHeader"] {visibility: visible !important; background-color: transparent !important;} .block-container {padding-top: 2rem !important;}</style>""", unsafe_allow_html=True)

verificar_acesso()

# ==============================================================================
# 3. LÓGICA DE BANCO DE DADOS E RASTRO (GOOGLE SHEETS)
# ==============================================================================
@st.cache_resource
def conectar_gsheets():
    try:
        scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
        credentials = Credentials.from_service_account_info(st.secrets["gcp_service_account"], scopes=scope)
        client = gspread.authorize(credentials)
        return client
    except: return None

# --- FUNÇÃO NOVA: SALVAR RASTRO ---
def salvar_rastro_validacao(aluno_nome, tipo_atividade, componente_tema, tema_objetivo):
    """
    Salva na aba 'Rastro_Validacao' da planilha Omnisfera_Dados.
    Campos: Data, Aluno, Tipo, Componente, Tema.
    """
    try:
        client = conectar_gsheets()
        if not client: return
        
        sheet = client.open("Omnisfera_Dados")
        try:
            worksheet = sheet.worksheet("Rastro_Validacao")
        except:
            # Cria a aba se não existir
            worksheet = sheet.add_worksheet(title="Rastro_Validacao", rows=1000, cols=5)
            worksheet.append_row(["Data", "Aluno", "Tipo Atividade", "Componente/Campo", "Tema/Objetivo"])
            
        data_hoje = date.today().strftime("%d/%m/%Y")
        worksheet.append_row([data_hoje, aluno_nome, tipo_atividade, componente_tema, tema_objetivo])
        st.toast("✅ Atividade validada e salva no histórico!", icon="💾")
    except Exception as e:
        st.error(f"Erro ao salvar rastro: {e}")

# Carrega banco de estudantes da memória ou local
if 'banco_estudantes' not in st.session_state:
    if os.path.exists("banco_alunos.json"):
        try:
            with open("banco_alunos.json", "r", encoding="utf-8") as f:
                st.session_state.banco_estudantes = json.load(f)
        except: st.session_state.banco_estudantes = []
    else:
        st.session_state.banco_estudantes = []

# ==============================================================================
# 4. FUNÇÕES DE UTILIDADE (DOCX, PDF, IMAGEM)
# ==============================================================================
def get_img_tag(file_path, width):
    if os.path.exists(file_path):
        with open(file_path, "rb") as f:
            data = base64.b64encode(f.read()).decode("utf-8")
        return f'<img src="data:image/png;base64,{data}" width="{width}">'
    return "🚀"

def extrair_dados_docx(uploaded_file):
    uploaded_file.seek(0); imagens = []; texto = ""
    try:
        doc = Document(uploaded_file)
        texto = "\n".join([p.text for p in doc.paragraphs if p.text.strip() != ""])
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                img_data = rel.target_part.blob
                if len(img_data) > 1024: imagens.append(img_data)
    except: pass
    return texto, imagens

def sanitizar_imagem(image_bytes):
    try:
        img = Image.open(BytesIO(image_bytes)).convert("RGB")
        out = BytesIO()
        img.save(out, format="JPEG", quality=90)
        return out.getvalue()
    except: return None

def baixar_imagem_url(url):
    try:
        resp = requests.get(url, timeout=10)
        if resp.status_code == 200: return BytesIO(resp.content)
    except: pass
    return None

def buscar_imagem_unsplash(query, access_key):
    if not access_key: return None
    url = f"https://api.unsplash.com/search/photos?query={query}&per_page=1&client_id={access_key}&lang=pt"
    try:
        resp = requests.get(url, timeout=5)
        data = resp.json()
        if data.get('results'):
            return data['results'][0]['urls']['regular']
    except: pass
    return None

def garantir_tag_imagem(texto):
    if "[[IMG" not in texto.upper() and "[[GEN_IMG" not in texto.upper():
        match = re.search(r'(\n|\. )', texto)
        if match:
            pos = match.end()
            return texto[:pos] + "\n\n[[IMG_1]]\n\n" + texto[pos:]
        return texto + "\n\n[[IMG_1]]"
    return texto

def construir_docx_final(texto_ia, aluno, materia, mapa_imgs, img_dalle_url, tipo_atv, sem_cabecalho=False):
    doc = Document(); style = doc.styles['Normal']; style.font.name = 'Arial'; style.font.size = Pt(12)
    if not sem_cabecalho:
        doc.add_heading(f'{tipo_atv.upper()} ADAPTADA - {materia.upper()}', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Estudante: {aluno['nome']}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("_"*50)
        doc.add_heading('Atividades', level=2)
    linhas = texto_ia.split('\n')
    for linha in linhas:
        tag_match = re.search(r'\[\[(IMG|GEN_IMG).*?(\d+)\]\]', linha, re.IGNORECASE)
        if tag_match:
            partes = re.split(r'(\[\[(?:IMG|GEN_IMG).*?\d+\]\])', linha, flags=re.IGNORECASE)
            for parte in partes:
                sub_match = re.search(r'(\d+)', parte)
                if ("IMG" in parte.upper() or "GEN_IMG" in parte.upper()) and sub_match:
                    num = int(sub_match.group(1))
                    img_bytes = mapa_imgs.get(num)
                    if not img_bytes and len(mapa_imgs) == 1: img_bytes = list(mapa_imgs.values())[0]
                    if img_bytes:
                        try:
                            p = doc.add_paragraph()
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            r = p.add_run()
                            r.add_picture(BytesIO(img_bytes), width=Inches(4.5))
                        except: pass
                elif parte.strip(): doc.add_paragraph(parte.strip())
        else:
            if linha.strip(): doc.add_paragraph(linha.strip())
    buffer = BytesIO(); doc.save(buffer); buffer.seek(0)
    return buffer

def criar_docx_simples(texto, titulo="Documento"):
    doc = Document()
    doc.add_heading(titulo, 0)
    for para in texto.split('\n'):
        if para.strip(): doc.add_paragraph(para.strip())
    buffer = BytesIO(); doc.save(buffer); buffer.seek(0)
    return buffer

def criar_pdf_generico(texto):
    pdf = FPDF(); pdf.add_page(); pdf.set_font("Arial", size=12)
    texto_safe = texto.encode('latin-1', 'replace').decode('latin-1')
    pdf.multi_cell(0, 10, texto_safe)
    return pdf.output(dest='S').encode('latin-1')

# --- IA FUNCTIONS ---
def gerar_imagem_inteligente(api_key, prompt, unsplash_key=None, feedback_anterior="", prioridade="IA"):
    client = OpenAI(api_key=api_key)
    prompt_final = f"{prompt}. Adjustment requested: {feedback_anterior}" if feedback_anterior else prompt
    if prioridade == "BANCO" and unsplash_key:
        url_banco = buscar_imagem_unsplash(prompt.split('.')[0], unsplash_key)
        if url_banco: return url_banco
    try:
        didactic_prompt = f"Educational textbook illustration, clean flat vector style, white background. CRITICAL RULE: STRICTLY NO TEXT. {prompt_final}"
        resp = client.images.generate(model="dall-e-3", prompt=didactic_prompt, size="1024x1024", quality="standard", n=1)
        return resp.data[0].url
    except: return None

def gerar_pictograma_caa(api_key, conceito, feedback_anterior=""):
    client = OpenAI(api_key=api_key)
    prompt_caa = f"Create a COMMUNICATION SYMBOL (AAC/PECS) for: '{conceito}'. White background, thick black outlines. MUTE IMAGE. NO TEXT."
    try:
        resp = client.images.generate(model="dall-e-3", prompt=prompt_caa, size="1024x1024", quality="standard", n=1)
        return resp.data[0].url
    except: return None

def gerar_experiencia_ei_bncc(api_key, aluno, campo_exp, objetivo, feedback_anterior=""):
    client = OpenAI(api_key=api_key)
    ajuste_prompt = f"AJUSTE: {feedback_anterior}." if feedback_anterior else ""
    prompt = f"""ATUAR COMO: Especialista EI BNCC. ALUNO: {aluno['nome']}. HIPERFOCO: {aluno.get('hiperfoco')}. PEI: {aluno.get('ia_sugestao', '')[:600]}. CRIAR EXPERIÊNCIA LÚDICA Campo: "{campo_exp}". Objetivo: {objetivo}. {ajuste_prompt} SAÍDA MARKDOWN."""
    try:
        resp = client.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "user", "content": prompt}], temperature=0.7)
        return resp.choices[0].message.content
    except Exception as e: return str(e)

def adaptar_conteudo_docx(api_key, aluno, texto, materia, tema, tipo_atv, remover_resp, questoes_mapeadas, modo_profundo=False):
    client = OpenAI(api_key=api_key)
    lista_q = ", ".join([str(n) for n in questoes_mapeadas])
    style = "Seja didático." if not modo_profundo else "Use Cadeia de Pensamento."
    prompt = f"""ESPECIALISTA DUA. {style} 1. ANALISE PEI: {aluno.get('ia_sugestao', '')[:1000]}. 2. ADAPTE usando hiperfoco ({aluno.get('hiperfoco')}). IMAGENS OBRIGATÓRIAS: {lista_q}. SAÍDA: [ANÁLISE PEDAGÓGICA] ... ---DIVISOR--- [ATIVIDADE] ... CONTEXTO: {materia} | {tema}. TEXTO: {texto}"""
    try:
        resp = client.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "user", "content": prompt}], temperature=0.7 if modo_profundo else 0.4)
        full_text = resp.choices[0].message.content
        if "---DIVISOR---" in full_text:
            parts = full_text.split("---DIVISOR---")
            return parts[0].replace("[ANÁLISE PEDAGÓGICA]", "").strip(), parts[1].replace("[ATIVIDADE]", "").strip()
        return "Análise indisponível.", full_text
    except Exception as e: return str(e), ""

def adaptar_conteudo_imagem(api_key, aluno, imagem_bytes, materia, tema, tipo_atv, livro_professor, modo_profundo=False):
    client = OpenAI(api_key=api_key)
    b64 = base64.b64encode(imagem_bytes).decode('utf-8')
    instrucao = "Remova gabarito." if livro_professor else ""
    prompt = f"""OCR e Adaptação. {instrucao} Adapte para {aluno['nome']} (PEI: {aluno.get('ia_sugestao', '')[:800]}). Hiperfoco: {aluno.get('hiperfoco')}. Insira tag [[IMG_1]]. SAÍDA: [ANÁLISE PEDAGÓGICA] ... ---DIVISOR--- [ATIVIDADE] ..."""
    msgs = [{"role": "user", "content": [{"type": "text", "text": prompt}, {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}}]}]
    try:
        resp = client.chat.completions.create(model="gpt-4o-mini", messages=msgs, temperature=0.7 if modo_profundo else 0.4)
        full = resp.choices[0].message.content
        if "---DIVISOR---" in full:
            pts = full.split("---DIVISOR---")
            return pts[0].replace("[ANÁLISE PEDAGÓGICA]", "").strip(), garantir_tag_imagem(pts[1].replace("[ATIVIDADE]", "").strip())
        return "Análise.", full
    except Exception as e: return str(e), ""

def criar_profissional(api_key, aluno, materia, objeto, qtd, tipo_q, qtd_imgs, verbos_bloom=None, modo_profundo=False):
    client = OpenAI(api_key=api_key)
    img_instr = f"Incluir imagens em {qtd_imgs} questões (use [[GEN_IMG: termo]])." if qtd_imgs > 0 else "Sem imagens."
    prompt = f"""Crie prova {materia} ({objeto}). QTD: {qtd} ({tipo_q}). Hiperfoco: {aluno.get('hiperfoco')}. {img_instr}. SAÍDA: [ANÁLISE PEDAGÓGICA] ... ---DIVISOR--- [ATIVIDADE] ..."""
    try:
        resp = client.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "user", "content": prompt}], temperature=0.8 if modo_profundo else 0.6)
        full = resp.choices[0].message.content
        if "---DIVISOR---" in full:
            pts = full.split("---DIVISOR---")
            return pts[0].replace("[ANÁLISE PEDAGÓGICA]", "").strip(), pts[1].replace("[ATIVIDADE]", "").strip()
        return "Análise.", full
    except Exception as e: return str(e), ""

def gerar_roteiro_aula(api_key, aluno, materia, assunto, feedback_anterior=""):
    client = OpenAI(api_key=api_key)
    prompt = f"Roteiro de Aula para {aluno['nome']}. Matéria: {materia}. Assunto: {assunto}. Hiperfoco: {aluno.get('hiperfoco')}. Estrutura: Objetivo, Estratégia, Prática, Avaliação."
    try:
        resp = client.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "user", "content": prompt}], temperature=0.7)
        return resp.choices[0].message.content
    except Exception as e: return str(e)

def gerar_quebra_gelo_profundo(api_key, aluno, materia, assunto, hiperfoco, tema_turma_extra=""):
    client = OpenAI(api_key=api_key)
    prompt = f"3 sugestões de Papo de Mestre (Quebra-gelo) conectando {assunto} com hiperfoco {hiperfoco}."
    try:
        resp = client.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "user", "content": prompt}], temperature=0.8)
        return resp.choices[0].message.content
    except Exception as e: return str(e)

def gerar_dinamica_inclusiva(api_key, aluno, materia, assunto, qtd_alunos, caracteristicas_turma, feedback_anterior=""):
    client = OpenAI(api_key=api_key)
    prompt = f"Dinâmica Inclusiva. {materia}: {assunto}. {qtd_alunos} alunos. Foco em {aluno['nome']} ({aluno.get('hiperfoco')})."
    try:
        resp = client.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "user", "content": prompt}], temperature=0.7)
        return resp.choices[0].message.content
    except Exception as e: return str(e)

def gerar_plano_aula_bncc(api_key, materia, assunto, metodologia, tecnica, qtd_alunos, recursos):
    client = OpenAI(api_key=api_key)
    prompt = f"Plano de Aula BNCC. {materia}: {assunto}. Metodologia: {metodologia} ({tecnica}). {qtd_alunos} alunos. Recursos: {recursos}. Saída: Habilidades, Objetivos, Passo a Passo."
    try:
        resp = client.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "user", "content": prompt}], temperature=0.7)
        return resp.choices[0].message.content
    except Exception as e: return str(e)
    # ==============================================================================
# INTERFACE
# ==============================================================================
with st.sidebar:
    if 'OPENAI_API_KEY' in st.secrets: api_key = st.secrets['OPENAI_API_KEY']; st.success("✅ OpenAI OK")
    else: api_key = st.text_input("Chave OpenAI:", type="password")
    st.markdown("---")
    if 'UNSPLASH_ACCESS_KEY' in st.secrets: unsplash_key = st.secrets['UNSPLASH_ACCESS_KEY']; st.success("✅ Unsplash OK")
    else: unsplash_key = st.text_input("Chave Unsplash:", type="password")
    st.markdown("---")
    if st.button("🏠 Voltar para Home", use_container_width=True): st.switch_page("Home.py")
    if st.button("🧹 Limpar Tudo", type="secondary"):
        for key in list(st.session_state.keys()):
            if key not in ['banco_estudantes', 'OPENAI_API_KEY', 'UNSPLASH_ACCESS_KEY', 'autenticado']: del st.session_state[key]
        st.rerun()

img_hub_html = get_img_tag("hub.png", "220") 
st.markdown(f"""
    <div class="header-hub">
        <div style="flex-shrink: 0;">{img_hub_html}</div>
        <div style="flex-grow: 1; text-align: center;"><p style="margin:0; color:#2C5282; font-size: 1.3rem; font-weight: 700;">Adaptação de Materiais & Criação</p></div>
    </div>
""", unsafe_allow_html=True)

if not st.session_state.banco_estudantes:
    st.warning("⚠️ Nenhum aluno encontrado. Cadastre no módulo PEI primeiro.")
    st.stop()

lista = [a['nome'] for a in st.session_state.banco_estudantes]
nome_aluno = st.selectbox("📂 Selecione o Estudante:", lista)
aluno = next(a for a in st.session_state.banco_estudantes if a['nome'] == nome_aluno)

serie_aluno = aluno.get('serie', '').lower()
is_ei = "infantil" in serie_aluno or "creche" in serie_aluno or "pré" in serie_aluno

st.markdown(f"""
    <div class="student-header">
        <div class="student-info-item"><div class="student-label">Nome</div><div class="student-value">{aluno.get('nome')}</div></div>
        <div class="student-info-item"><div class="student-label">Série</div><div class="student-value">{aluno.get('serie', '-')}</div></div>
        <div class="student-info-item"><div class="student-label">Hiperfoco</div><div class="student-value">{aluno.get('hiperfoco', '-')}</div></div>
    </div>
""", unsafe_allow_html=True)

if 'res_scene_url' not in st.session_state: st.session_state.res_scene_url = None
if 'valid_scene' not in st.session_state: st.session_state.valid_scene = False
if 'res_caa_url' not in st.session_state: st.session_state.res_caa_url = None
if 'valid_caa' not in st.session_state: st.session_state.valid_caa = False

if is_ei:
    st.info("🧸 **Modo Educação Infantil Ativado:** Foco em Experiências, BNCC e Brincar.")
    tabs = st.tabs(["🧸 Criar Experiência (BNCC)", "🎨 Estúdio Visual & CAA", "📝 Rotina & AVD", "🤝 Inclusão no Brincar"])
    
    with tabs[0]:
        st.markdown("""<div class="pedagogia-box"><div class="pedagogia-title"><i class="ri-lightbulb-line"></i> Pedagogia do Brincar (BNCC)</div>Experiências de aprendizagem intencionais.</div>""", unsafe_allow_html=True)
        col_ei1, col_ei2 = st.columns(2)
        campo_exp = col_ei1.selectbox("Campo de Experiência (BNCC)", ["O eu, o outro e o nós", "Corpo, gestos e movimentos", "Traços, sons, cores e formas", "Escuta, fala, pensamento e imaginação", "Espaços, tempos, quantidades..."])
        obj_aprendizagem = col_ei2.text_input("Objetivo de Aprendizagem:", placeholder="Ex: Compartilhar brinquedos...")
        if 'res_ei_exp' not in st.session_state: st.session_state.res_ei_exp = None
        if 'valid_ei_exp' not in st.session_state: st.session_state.valid_ei_exp = False
        if st.button("✨ GERAR EXPERIÊNCIA LÚDICA", type="primary"):
            with st.spinner("Criando vivência..."):
                st.session_state.res_ei_exp = gerar_experiencia_ei_bncc(api_key, aluno, campo_exp=campo_exp, objetivo=obj_aprendizagem)
                st.session_state.valid_ei_exp = False
        if st.session_state.res_ei_exp:
            if st.session_state.valid_ei_exp:
                st.markdown("<div class='validado-box'>✅ EXPERIÊNCIA APROVADA!</div>", unsafe_allow_html=True)
                st.markdown(st.session_state.res_ei_exp)
            else:
                st.markdown(st.session_state.res_ei_exp)
                st.write("---")
                c_val, c_ref = st.columns([1, 3])
                if c_val.button("✅ Validar Experiência"): 
                    st.session_state.valid_ei_exp = True
                    # SALVA O RASTRO EI
                    salvar_rastro_validacao(aluno['nome'], "Experiência EI", campo_exp, obj_aprendizagem)
                    st.rerun()
                with c_ref.expander("🔄 Refazer"):
                    feedback_ei = st.text_input("Ajuste:", placeholder="Ex: Simplificar...")
                    if st.button("Refazer com Ajustes"):
                        with st.spinner("Reescrevendo..."):
                            st.session_state.res_ei_exp = gerar_experiencia_ei_bncc(api_key, aluno, campo_exp, obj_aprendizagem, feedback_anterior=feedback_ei)
                            st.rerun()

    with tabs[1]:
        st.markdown("""<div class="pedagogia-box"><div class="pedagogia-title"><i class="ri-eye-line"></i> Apoio Visual & Comunicação</div>Cenas e Pictogramas.</div>""", unsafe_allow_html=True)
        col_scene, col_caa = st.columns(2)
        with col_scene:
            st.markdown("#### 🖼️ Ilustração de Cena")
            desc_m = st.text_area("Descreva a cena:", height=100, key="vdm_ei")
            if st.button("🎨 Gerar Cena", key="btn_cena_ei"):
                with st.spinner("Desenhando..."):
                    st.session_state.res_scene_url = gerar_imagem_inteligente(api_key, f"{desc_m}. Child education.", unsplash_key)
                    st.session_state.valid_scene = False
            if st.session_state.res_scene_url:
                st.image(st.session_state.res_scene_url)
                if st.session_state.valid_scene: st.success("Validada!")
                else:
                    if st.button("✅ Validar", key="val_sc_ei"): 
                        st.session_state.valid_scene = True
                        salvar_rastro_validacao(aluno['nome'], "Ilustração Cena", "Artes Visuais", desc_m)
                        st.rerun()
        with col_caa:
            st.markdown("#### 🗣️ Símbolo CAA")
            palavra_chave = st.text_input("Conceito:", key="caa_input")
            if st.button("🧩 Gerar Pictograma", key="btn_caa"):
                with st.spinner("Criando..."):
                    st.session_state.res_caa_url = gerar_pictograma_caa(api_key, palavra_chave)
                    st.session_state.valid_caa = False
            if st.session_state.res_caa_url:
                st.image(st.session_state.res_caa_url, width=300)
                if st.session_state.valid_caa: st.success("Validado!")
                else:
                    if st.button("✅ Validar", key="val_caa_ei"): 
                        st.session_state.valid_caa = True
                        salvar_rastro_validacao(aluno['nome'], "Pictograma CAA", "Comunicação", palavra_chave)
                        st.rerun()

    with tabs[2]:
        st.markdown("""<div class="pedagogia-box"><div class="pedagogia-title"><i class="ri-calendar-check-line"></i> Rotina</div>Adaptação de rotina.</div>""", unsafe_allow_html=True)
        rotina_detalhada = st.text_area("Rotina:", height=200)
        topico_foco = st.text_input("Foco:", placeholder="Ex: Transição")
        if 'res_ei_rotina' not in st.session_state: st.session_state.res_ei_rotina = None
        if 'valid_ei_rotina' not in st.session_state: st.session_state.valid_ei_rotina = False
        if st.button("📝 ADAPTAR ROTINA", type="primary"):
            with st.spinner("Analisando..."):
                prompt_rotina = f"Rotina EI:\n{rotina_detalhada}\nFoco: {topico_foco}"
                st.session_state.res_ei_rotina = gerar_roteiro_aula(api_key, aluno, "Geral", "Rotina", feedback_anterior=prompt_rotina)
                st.session_state.valid_ei_rotina = False
        if st.session_state.res_ei_rotina:
            if st.session_state.valid_ei_rotina:
                st.markdown("<div class='validado-box'>✅ VALIDADA!</div>", unsafe_allow_html=True)
                st.markdown(st.session_state.res_ei_rotina)
            else:
                st.markdown(st.session_state.res_ei_rotina)
                if st.button("✅ Validar Rotina"): 
                    st.session_state.valid_ei_rotina = True
                    salvar_rastro_validacao(aluno['nome'], "Adaptação Rotina", "Rotina", topico_foco)
                    st.rerun()

    with tabs[3]:
        st.markdown("""<div class="pedagogia-box"><div class="pedagogia-title"><i class="ri-group-line"></i> Mediação Social</div>Pontes de interação.</div>""", unsafe_allow_html=True)
        tema_d = st.text_input("Tema:", key="dina_ei")
        if 'res_ei_dina' not in st.session_state: st.session_state.res_ei_dina = None
        if 'valid_ei_dina' not in st.session_state: st.session_state.valid_ei_dina = False
        if st.button("🤝 GERAR DINÂMICA", type="primary"): 
            with st.spinner("Criando..."):
                st.session_state.res_ei_dina = gerar_dinamica_inclusiva(api_key, aluno, "EI", tema_d, "grupo", "Crianças")
                st.session_state.valid_ei_dina = False
        if st.session_state.res_ei_dina:
            if st.session_state.valid_ei_dina:
                st.markdown("<div class='validado-box'>✅ VALIDADA!</div>", unsafe_allow_html=True)
                st.markdown(st.session_state.res_ei_dina)
            else:
                st.markdown(st.session_state.res_ei_dina)
                if st.button("✅ Validar Dinâmica"): 
                    st.session_state.valid_ei_dina = True
                    salvar_rastro_validacao(aluno['nome'], "Dinâmica Social", "EI", tema_d)
                    st.rerun()

else:
    tabs = st.tabs(["📄 Adaptar Prova", "✂️ Adaptar Atividade", "✨ Criar do Zero", "🎨 Estúdio Visual & CAA", "📝 Roteiro Individual", "🗣️ Papo de Mestre", "🤝 Dinâmica Inclusiva", "📅 Plano de Aula DUA"])

    with tabs[0]:
        st.markdown("""<div class="pedagogia-box"><div class="pedagogia-title"><i class="ri-file-edit-line"></i> Adaptação DUA</div>Provas acessíveis.</div>""", unsafe_allow_html=True)
        c1, c2, c3 = st.columns(3)
        materia_d = c1.selectbox("Matéria", ["Matemática", "Português", "Ciências", "História", "Geografia", "Artes", "Inglês"], key="dm")
        tema_d = c2.text_input("Tema", key="dt")
        tipo_d = c3.selectbox("Tipo", ["Prova", "Tarefa"], key="dtp")
        arquivo_d = st.file_uploader("Upload DOCX", type=["docx"], key="fd")
        if 'docx_imgs' not in st.session_state: st.session_state.docx_imgs = []
        if 'docx_txt' not in st.session_state: st.session_state.docx_txt = None
        if arquivo_d and arquivo_d.file_id != st.session_state.get('last_d'):
            st.session_state.last_d = arquivo_d.file_id
            txt, imgs = extrair_dados_docx(arquivo_d)
            st.session_state.docx_txt = txt; st.session_state.docx_imgs = imgs
            st.success(f"{len(imgs)} imagens.")
        map_d = {}; qs_d = []
        if st.session_state.docx_imgs:
            cols = st.columns(3)
            for i, img in enumerate(st.session_state.docx_imgs):
                with cols[i % 3]:
                    st.image(img, width=80)
                    q = st.number_input(f"Questão:", 0, 50, key=f"dq_{i}")
                    if q > 0: map_d[int(q)] = img; qs_d.append(int(q))
        if st.button("🚀 ADAPTAR PROVA", type="primary", key="btn_d"):
            if not st.session_state.docx_txt: st.warning("Envie arquivo."); st.stop()
            with st.spinner("Adaptando..."):
                rac, txt = adaptar_conteudo_docx(api_key, aluno, st.session_state.docx_txt, materia_d, tema_d, tipo_d, True, qs_d)
                st.session_state['res_docx'] = {'rac': rac, 'txt': txt, 'map': map_d, 'valid': False}
                st.rerun()
        if 'res_docx' in st.session_state:
            res = st.session_state['res_docx']
            if res.get('valid'):
                st.markdown("<div class='validado-box'>✅ VALIDADO!</div>", unsafe_allow_html=True)
            else:
                col_v, col_r = st.columns([1, 1])
                if col_v.button("✅ Validar", key="val_d"): 
                    st.session_state['res_docx']['valid'] = True
                    salvar_rastro_validacao(aluno['nome'], "Adaptação Prova", materia_d, tema_d)
                    st.rerun()
                if col_r.button("🧠 Refazer (+Profundo)", key="redo_d"):
                    with st.spinner("Refazendo..."):
                        rac, txt = adaptar_conteudo_docx(api_key, aluno, st.session_state.docx_txt, materia_d, tema_d, tipo_d, True, qs_d, modo_profundo=True)
                        st.session_state['res_docx'] = {'rac': rac, 'txt': txt, 'map': map_d, 'valid': False}
                        st.rerun()
            st.markdown(f"<div class='analise-box'>{res['rac']}</div>", unsafe_allow_html=True)
            docx = construir_docx_final(res['txt'], aluno, materia_d, res['map'], None, tipo_d)
            st.download_button("📥 BAIXAR DOCX", docx, "Prova_Adaptada.docx", "primary")

    with tabs[1]:
        st.markdown("""<div class="pedagogia-box"><div class="pedagogia-title"><i class="ri-scissors-cut-line"></i> OCR</div>Adaptação de foto.</div>""", unsafe_allow_html=True)
        c1, c2, c3 = st.columns(3)
        materia_i = c1.selectbox("Matéria", ["Matemática", "Português", "Ciências", "História"], key="im")
        tema_i = c2.text_input("Tema", key="it")
        tipo_i = c3.selectbox("Tipo", ["Atividade", "Tarefa"], key="itp")
        arquivo_i = st.file_uploader("Upload", type=["png","jpg"], key="fi")
        if 'img_raw' not in st.session_state: st.session_state.img_raw = None
        if arquivo_i and arquivo_i.file_id != st.session_state.get('last_i'):
            st.session_state.last_i = arquivo_i.file_id
            st.session_state.img_raw = sanitizar_imagem(arquivo_i.getvalue())
        cropped_res = None
        if st.session_state.img_raw:
            img_pil = Image.open(BytesIO(st.session_state.img_raw))
            img_pil.thumbnail((800, 800))
            cropped_res = st_cropper(img_pil, realtime_update=True, box_color='#FF0000', aspect_ratio=None, key="crop_i")
        if st.button("🚀 ADAPTAR ATIVIDADE", type="primary", key="btn_i"):
            if not st.session_state.img_raw: st.warning("Envie imagem."); st.stop()
            with st.spinner("Adaptando..."):
                buf_c = BytesIO(); cropped_res.convert('RGB').save(buf_c, format="JPEG", quality=90); img_bytes = buf_c.getvalue()
                rac, txt = adaptar_conteudo_imagem(api_key, aluno, img_bytes, materia_i, tema_i, tipo_i, False)
                st.session_state['res_img'] = {'rac': rac, 'txt': txt, 'map': {1: img_bytes}, 'valid': False}
                st.rerun()
        if 'res_img' in st.session_state:
            res = st.session_state['res_img']
            if res.get('valid'): st.markdown("<div class='validado-box'>✅ VALIDADO!</div>", unsafe_allow_html=True)
            else:
                col_v, col_r = st.columns([1, 1])
                if col_v.button("✅ Validar", key="val_i"): 
                    st.session_state['res_img']['valid'] = True
                    salvar_rastro_validacao(aluno['nome'], "Adaptação OCR", materia_i, tema_i)
                    st.rerun()
                if col_r.button("🧠 Refazer", key="redo_i"):
                    with st.spinner("Refazendo..."):
                        img_bytes = res['map'][1]
                        rac, txt = adaptar_conteudo_imagem(api_key, aluno, img_bytes, materia_i, tema_i, tipo_i, False, modo_profundo=True)
                        st.session_state['res_img'] = {'rac': rac, 'txt': txt, 'map': {1: img_bytes}, 'valid': False}
                        st.rerun()
            st.markdown(f"<div class='analise-box'>{res['rac']}</div>", unsafe_allow_html=True)
            docx = construir_docx_final(res['txt'], aluno, materia_i, res['map'], None, tipo_i)
            st.download_button("📥 BAIXAR DOCX", docx, "Atividade.docx", "primary")

    with tabs[2]:
        st.markdown("""<div class="pedagogia-box"><div class="pedagogia-title"><i class="ri-magic-line"></i> Criação</div>Do zero.</div>""", unsafe_allow_html=True)
        cc1, cc2 = st.columns(2)
        mat_c = cc1.selectbox("Componente", ["Matemática", "Português", "Ciências"], key="cm")
        obj_c = cc2.text_input("Assunto", key="co")
        cc3, cc4 = st.columns(2)
        qtd_c = cc3.slider("Qtd Questões", 1, 10, 5, key="cq")
        tipo_quest = cc4.selectbox("Tipo", ["Objetiva", "Discursiva"], key="ctq")
        usar_img = st.checkbox("Imagens?", value=True)
        qtd_img_sel = st.slider("Qtd Imagens", 0, qtd_c, int(qtd_c/2), disabled=not usar_img)
        if st.button("✨ CRIAR", type="primary", key="btn_c"):
            with st.spinner("Elaborando..."):
                qtd_final = qtd_img_sel if usar_img else 0
                rac, txt = criar_profissional(api_key, aluno, mat_c, obj_c, qtd_c, tipo_quest, qtd_final)
                novo_map = {}; count = 0
                tags = re.findall(r'\[\[GEN_IMG: (.*?)\]\]', txt)
                for p in tags:
                    count += 1
                    url = gerar_imagem_inteligente(api_key, p, unsplash_key, prioridade="BANCO")
                    if url:
                        io = baixar_imagem_url(url)
                        if io: novo_map[count] = io.getvalue()
                txt_fin = txt
                for i in range(1, count + 1): txt_fin = re.sub(r'\[\[GEN_IMG: .*?\]\]', f"[[IMG_G{i}]]", txt_fin, count=1)
                st.session_state['res_create'] = {'rac': rac, 'txt': txt_fin, 'map': novo_map, 'valid': False}
                st.rerun()
        if 'res_create' in st.session_state:
            res = st.session_state['res_create']
            if res.get('valid'): st.markdown("<div class='validado-box'>✅ VALIDADO!</div>", unsafe_allow_html=True)
            else:
                col_v, col_r = st.columns([1, 1])
                if col_v.button("✅ Validar", key="val_c"): 
                    st.session_state['res_create']['valid'] = True
                    salvar_rastro_validacao(aluno['nome'], "Criação DUA", mat_c, obj_c)
                    st.rerun()
                if col_r.button("🧠 Refazer", key="redo_c"):
                    with st.spinner("Refazendo..."):
                        qtd_final = qtd_img_sel if usar_img else 0
                        rac, txt = criar_profissional(api_key, aluno, mat_c, obj_c, qtd_c, tipo_quest, qtd_final, modo_profundo=True)
                        st.session_state['res_create']['rac'] = rac
                        st.session_state['res_create']['txt'] = txt
                        st.session_state['res_create']['valid'] = False
                        st.rerun()
            st.markdown(f"<div class='analise-box'>{res['rac']}</div>", unsafe_allow_html=True)
            docx = construir_docx_final(res['txt'], aluno, mat_c, {}, None, "Criada")
            st.download_button("📥 DOCX", docx, "Criada.docx", "primary")

    with tabs[3]:
        st.markdown("#### Estúdio Visual & CAA")
        c1, c2 = st.columns(2)
        with c1:
            desc_m = st.text_area("Cena:", key="vdm_pd")
            if st.button("Gerar Cena", key="btn_cn_pd"):
                st.session_state.res_scene_url = gerar_imagem_inteligente(api_key, desc_m)
                st.session_state.valid_scene = False
            if st.session_state.res_scene_url:
                st.image(st.session_state.res_scene_url)
                if st.session_state.valid_scene: st.success("Validado!")
                elif st.button("Validar Cena", key="val_cn_pd"): 
                    st.session_state.valid_scene = True
                    salvar_rastro_validacao(aluno['nome'], "Ilustração Cena", "Artes", desc_m)
                    st.rerun()
        with c2:
            conc = st.text_input("Pictograma:", key="caa_pd")
            if st.button("Gerar Picto", key="btn_caa_pd"):
                st.session_state.res_caa_url = gerar_pictograma_caa(api_key, conc)
                st.session_state.valid_caa = False
            if st.session_state.res_caa_url:
                st.image(st.session_state.res_caa_url)
                if st.session_state.valid_caa: st.success("Validado!")
                elif st.button("Validar Picto", key="val_caa_pd"): 
                    st.session_state.valid_caa = True
                    salvar_rastro_validacao(aluno['nome'], "Pictograma CAA", "Comunicação", conc)
                    st.rerun()

    with tabs[4]:
        st.markdown("#### Roteiro Individual")
        c1, c2 = st.columns(2)
        materia_rotina = c1.selectbox("Matéria", ["Matemática", "Português"], key="rot_mat")
        tema_rotina = c2.text_input("Tema", key="rot_tema")
        if st.button("Gerar Roteiro", key="btn_rot"):
            res = gerar_roteiro_aula(api_key, aluno, materia_rotina, tema_rotina)
            st.session_state['res_roteiro'] = res
            st.session_state['valid_roteiro'] = False
        if 'res_roteiro' in st.session_state:
            if st.session_state.get('valid_roteiro'): st.success("Validado!")
            else:
                st.markdown(st.session_state['res_roteiro'])
                if st.button("Validar Roteiro", key="val_rot"): 
                    st.session_state['valid_roteiro'] = True
                    salvar_rastro_validacao(aluno['nome'], "Roteiro Individual", materia_rotina, tema_rotina)
                    st.rerun()

    with tabs[5]:
        st.markdown("#### Papo de Mestre")
        c1, c2 = st.columns(2)
        materia_papo = c1.selectbox("Matéria", ["Matemática", "Português"], key="pap_mat")
        assunto_papo = c2.text_input("Assunto", key="pap_ass")
        if st.button("Gerar Papo", key="btn_pap"):
            res = gerar_quebra_gelo_profundo(api_key, aluno, materia_papo, assunto_papo, aluno.get('hiperfoco'))
            st.session_state['res_papo'] = res
            st.session_state['valid_papo'] = False
        if 'res_papo' in st.session_state:
            if st.session_state.get('valid_papo'): st.success("Validado!")
            else:
                st.markdown(st.session_state['res_papo'])
                if st.button("Validar Papo", key="val_pap"): 
                    st.session_state['valid_papo'] = True
                    salvar_rastro_validacao(aluno['nome'], "Papo de Mestre", materia_papo, assunto_papo)
                    st.rerun()

    with tabs[6]:
        st.markdown("#### Dinâmica Inclusiva")
        c1, c2 = st.columns(2)
        materia_din = c1.selectbox("Matéria", ["Matemática", "Português"], key="din_mat")
        assunto_din = c2.text_input("Assunto", key="din_ass")
        if st.button("Gerar Dinâmica", key="btn_din"):
            res = gerar_dinamica_inclusiva(api_key, aluno, materia_din, assunto_din, 25, "Geral")
            st.session_state['res_dinamica'] = res
            st.session_state['valid_dinamica'] = False
        if 'res_dinamica' in st.session_state:
            if st.session_state.get('valid_dinamica'): st.success("Validado!")
            else:
                st.markdown(st.session_state['res_dinamica'])
                if st.button("Validar Dinâmica", key="val_din"): 
                    st.session_state['valid_dinamica'] = True
                    salvar_rastro_validacao(aluno['nome'], "Dinâmica Inclusiva", materia_din, assunto_din)
                    st.rerun()

    with tabs[7]:
        st.markdown("#### Plano de Aula DUA")
        c1, c2 = st.columns(2)
        materia_plano = c1.selectbox("Matéria", ["Matemática", "Português"], key="pla_mat")
        assunto_plano = c2.text_input("Assunto", key="pla_ass")
        if st.button("Gerar Plano", key="btn_pla"):
            res = gerar_plano_aula_bncc(api_key, materia_plano, assunto_plano, "Ativa", "Rotação", 30, ["Projetor"])
            st.session_state['res_plano'] = res
            st.session_state['valid_plano'] = False
        if 'res_plano' in st.session_state:
            if st.session_state.get('valid_plano'): st.success("Validado!")
            else:
                st.markdown(st.session_state['res_plano'])
                if st.button("Validar Plano", key="val_pla"): 
                    st.session_state['valid_plano'] = True
                    salvar_rastro_validacao(aluno['nome'], "Plano de Aula DUA", materia_plano, assunto_plano)
                    st.rerun()
