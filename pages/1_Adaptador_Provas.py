import streamlit as st
from openai import OpenAI
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
from streamlit_cropper import st_cropper
import re
import requests
import json
import base64

# --- 1. CONFIGURAÇÃO ---
st.set_page_config(page_title="Adaptador 360º | V11.0", page_icon="🧩", layout="wide")

# --- 2. BANCO DE DADOS ---
ARQUIVO_DB = "banco_alunos.json"

def carregar_banco():
    if os.path.exists(ARQUIVO_DB):
        try:
            with open(ARQUIVO_DB, "r", encoding="utf-8") as f:
                return json.load(f)
        except: return []
    return []

if 'banco_estudantes' not in st.session_state or not st.session_state.banco_estudantes:
    st.session_state.banco_estudantes = carregar_banco()

# --- 3. ESTILO VISUAL ---
st.markdown("""
    <style>
    html, body, [class*="css"] { font-family: 'Nunito', sans-serif; color: #2D3748; }
    .header-clean { background: white; padding: 25px; border-radius: 16px; border: 1px solid #EDF2F7; margin-bottom: 20px; display: flex; gap: 20px; align-items: center; }
    .student-header { background-color: #EBF8FF; border: 1px solid #BEE3F8; border-radius: 12px; padding: 15px 25px; margin-bottom: 25px; display: flex; justify-content: space-between; align-items: center; }
    .student-label { font-size: 0.85rem; color: #718096; font-weight: 700; text-transform: uppercase; }
    .student-value { font-size: 1.1rem; color: #2C5282; font-weight: 800; }
    .crop-instruction { background: #EBF8FF; border-left: 4px solid #3182CE; padding: 15px; color: #2C5282; border-radius: 4px; margin-bottom: 10px; }
    
    .stTabs [data-baseweb="tab-list"] { gap: 8px; }
    .stTabs [data-baseweb="tab"] { border-radius: 4px; padding: 10px 20px; background-color: white; border: 1px solid #E2E8F0; }
    .stTabs [aria-selected="true"] { background-color: #3182CE !important; color: white !important; }

    div[data-testid="column"] .stButton button[kind="primary"] { border-radius: 12px !important; height: 50px; width: 100%; background-color: #FF6B6B !important; color: white !important; font-weight: 800 !important; }
    div[data-testid="column"] .stButton button[kind="secondary"] { border-radius: 12px !important; height: 50px; width: 100%; background-color: white !important; color: #718096 !important; border: 2px solid #CBD5E0 !important; }
    </style>
""", unsafe_allow_html=True)

# --- 4. FUNÇÕES DE ARQUIVO ---
def extrair_dados_docx(uploaded_file):
    uploaded_file.seek(0); imagens = []; texto = ""
    try:
        doc = Document(uploaded_file)
        texto = "\n".join([p.text for p in doc.paragraphs if p.text.strip() != ""])
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                img_data = rel.target_part.blob
                if len(img_data) > 1024: imagens.append(img_data)
    except: pass
    return texto, imagens

def baixar_imagem_url(url):
    try:
        resp = requests.get(url, timeout=10)
        if resp.status_code == 200: return BytesIO(resp.content)
    except: pass
    return None

def construir_docx_final(texto_ia, aluno, materia, mapa_imgs, img_dalle_url, tipo_atv):
    doc = Document(); style = doc.styles['Normal']; style.font.name = 'Arial'; style.font.size = Pt(12)
    doc.add_heading(f'{tipo_atv.upper()} ADAPTADA - {materia.upper()}', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Estudante: {aluno['nome']}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("_"*50)

    if img_dalle_url:
        img_io = baixar_imagem_url(img_dalle_url)
        if img_io:
            doc.add_heading('Apoio Visual', level=3)
            doc.add_picture(img_io, width=Inches(4.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("")

    doc.add_heading('Atividades', level=2)
    
    partes = re.split(r'(\[\[IMG_[Q|G]\w+\]\])', texto_ia)
    
    for parte in partes:
        tag_match = re.search(r'\[\[IMG_(Q|G)(\w+)\]\]', parte)
        if tag_match:
            tipo, id_img = tag_match.groups()
            img_bytes = None
            if tipo == "Q": 
                try: num = int(id_img); img_bytes = mapa_imgs.get(num, mapa_imgs.get(0))
                except: pass
            elif tipo == "G": 
                img_bytes = mapa_imgs.get(f"G{id_img}")

            if img_bytes:
                try:
                    doc.add_picture(BytesIO(img_bytes), width=Inches(4.5))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph("") 
                except: pass
        elif parte.strip():
            doc.add_paragraph(parte.strip())
            
    buffer = BytesIO(); doc.save(buffer); buffer.seek(0)
    return buffer

# --- 5. INTELIGÊNCIA ARTIFICIAL ---
def gerar_dalle_prompt(api_key, prompt_text):
    client = OpenAI(api_key=api_key)
    try:
        resp = client.images.generate(model="dall-e-3", prompt=prompt_text + " Educational style, clear, autism-friendly, white background, no text.", size="1024x1024", quality="standard", n=1)
        return resp.data[0].url
    except: return None

# ADAPTAR DOCX (Prova)
def adaptar_conteudo_docx(api_key, aluno, texto, materia, tema, tipo_atv, remover_resp, questoes_mapeadas):
    client = OpenAI(api_key=api_key)
    lista_q = ", ".join([str(n) for n in questoes_mapeadas])
    
    prompt = f"""
    VOCÊ É UM ESPECIALISTA EM ADAPTAÇÃO CURRICULAR (DUA).
    
    MISSÃO:
    Adapte o conteúdo abaixo para {aluno['nome']}, considerando seu PEI e Hiperfoco.
    
    REGRAS DE OURO:
    1. **Incidência do Hiperfoco ({aluno.get('hiperfoco', 'Geral')}):** Aplique o tema em cerca de 20% a 30% das questões. Escolha aquelas onde o contexto se encaixa melhor. Nas outras, mantenha o contexto original ou simplifique.
    2. **Imagens:** O professor mapeou imagens nas questões: {lista_q}. INSIRA A TAG [[IMG_QX]] (onde X é o número) logo após o enunciado dessas questões.
    3. **Rigor:** Mantenha o objetivo pedagógico original. Não infantilize se não for necessário.
    
    SAÍDA OBRIGATÓRIA:
    [RACIONAL] (Explique suas escolhas)
    ---DIVISOR---
    [ATIVIDADE] (Texto final formatado)
    
    PEI: {aluno.get('ia_sugestao', '')[:1200]}
    {"REMOVA GABARITO." if remover_resp else ""}
    CONTEXTO: {materia} | {tema}
    TEXTO ORIGINAL: {texto}
    """
    try:
        resp = client.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "user", "content": prompt}], temperature=0.4)
        parts = resp.choices[0].message.content.split("---DIVISOR---")
        return (parts[0].strip(), parts[1].strip()) if len(parts)>1 else ("Adaptado.", resp.choices[0].message.content)
    except Exception as e: return str(e), ""

# ADAPTAR ATIVIDADE (Imagem/Recorte)
def adaptar_conteudo_imagem(api_key, aluno, imagem_bytes, materia, tema, tipo_atv, remover_resp):
    client = OpenAI(api_key=api_key)
    
    if isinstance(imagem_bytes, bytes):
        b64 = base64.b64encode(imagem_bytes).decode('utf-8')
    else: return "Erro formato img", ""

    prompt = f"""
    VOCÊ É UM LEITOR E ADAPTADOR DE ATIVIDADES.
    
    PASSO 1 (OCR): Leia atentamente todo o texto presente na imagem.
    PASSO 2 (ADAPTAÇÃO): Reescreva o conteúdo lido adaptando para o aluno.
    
    DIRETRIZES:
    - Use o PEI do aluno para ajustar a complexidade.
    - Se possível, faça uma leve conexão com o hiperfoco ({aluno.get('hiperfoco')}) para engajar (regra dos 30%).
    - Insira a tag [[IMG_Q1]] logo após o enunciado principal para manter a referência visual da figura original.
    
    SAÍDA: [RACIONAL] ---DIVISOR--- [ATIVIDADE]
    
    PEI: {aluno.get('ia_sugestao', '')[:1000]}
    {"REMOVA RESPOSTAS/GABARITO." if remover_resp else ""}
    CONTEXTO: {materia} | {tema}
    """
    
    msgs = [{"role": "user", "content": [{"type": "text", "text": prompt}, {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}}]}]

    try:
        resp = client.chat.completions.create(model="gpt-4o-mini", messages=msgs, temperature=0.4)
        parts = resp.choices[0].message.content.split("---DIVISOR---")
        return (parts[0].strip(), parts[1].strip()) if len(parts)>1 else ("Adaptado.", resp.choices[0].message.content)
    except Exception as e: return str(e), ""

# CRIAR DO ZERO (PROFISSIONAL)
def criar_profissional(api_key, aluno, materia, objeto, qtd, tipo_q):
    client = OpenAI(api_key=api_key)
    hiperfoco = aluno.get('hiperfoco', 'Geral')
    
    prompt = f"""
    VOCÊ É UM PROFESSOR SÊNIOR ELABORADOR DE ITENS.
    Sua tarefa é criar uma atividade de {materia} ({objeto}) com qualidade profissional.
    
    ALUNO: {aluno['nome']} | Série: {aluno.get('serie')}
    HIPERFOCO: {hiperfoco} (Use como contexto motivador em 30% das questões).
    
    REGRAS DE QUALIDADE:
    1. **Densidade e Contexto:** As questões não podem ser rasas. Devem ter um texto base ou situação-problema clara.
    2. **Rigor BNCC:** O conteúdo deve ser adequado à série, nem muito fácil, nem impossível.
    3. **Engajamento Inteligente:** Use o hiperfoco de forma orgânica (ex: calcular a área de uma construção do Minecraft), não forçada.
    4. **Imagens:** A cada 3 ou 4 questões, sugira uma imagem com a tag [[GEN_IMG: descrição detalhada]].
    
    SAÍDA:
    [RACIONAL PEDAGÓGICO]
    ---DIVISOR---
    [ATIVIDADE COMPLETA]
    """
    try:
        resp = client.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "user", "content": prompt}], temperature=0.6)
        full = resp.choices[0].message.content
        parts = full.split("---DIVISOR---")
        return (parts[0].strip(), parts[1].strip()) if len(parts)>1 else ("Criado.", full)
    except Exception as e: return str(e), ""

# CONTEXTUALIZAR
def gerar_contextualizacao(api_key, aluno, assunto, tema_extra=""):
    client = OpenAI(api_key=api_key)
    tema = tema_extra if tema_extra else aluno.get('hiperfoco', 'Geral')
    prompt = f"""
    CRIE UM ROTEIRO DE AULA CONTEXTUALIZADO.
    Assunto: {assunto}
    Público: {aluno['nome']} (Série: {aluno.get('serie')})
    Tema de Conexão: {tema}
    
    GERE:
    1. **Abertura (Quebra-Gelo):** Uma curiosidade ou pergunta sobre {tema} que introduz o assunto.
    2. **Explicação Principal:** O conceito de {assunto} explicado com analogias do universo de {tema}.
    3. **Estratégia de Engajamento:** Uma dinâmica rápida sugerida.
    """
    try:
        resp = client.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "user", "content": prompt}], temperature=0.7)
        return resp.choices[0].message.content
    except Exception as e: return str(e)

# SUGESTÃO DE IMAGEM VIA PEI
def sugerir_imagem_pei(api_key, aluno):
    client = OpenAI(api_key=api_key)
    prompt = f"""
    Analise o PEI deste aluno e seu hiperfoco.
    Descreva UM recurso visual (imagem) que seria perfeito para ajudar na regulação emocional ou foco dele AGORA.
    Ex: Uma prancha de rotina, um personagem fazendo sinal de silêncio, etc.
    Responda APENAS com a descrição visual em inglês para o DALL-E.
    
    PEI: {aluno.get('ia_sugestao', '')[:1000]}
    Hiperfoco: {aluno.get('hiperfoco')}
    """
    try:
        resp = client.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "user", "content": prompt}], temperature=0.7)
        return resp.choices[0].message.content
    except: return "Educational illustration"

# --- 6. INTERFACE ---
with st.sidebar:
    if 'OPENAI_API_KEY' in st.secrets: api_key = st.secrets['OPENAI_API_KEY']; st.success("✅ Conectado")
    else: api_key = st.text_input("Chave OpenAI:", type="password")
    st.markdown("---")
    if st.button("🗑️ Nova Sessão"):
        for k in list(st.session_state.keys()):
            if k not in ['banco_estudantes', 'OPENAI_API_KEY']: del st.session_state[k]
        st.rerun()

st.markdown("""<div class="header-clean"><div style="font-size:3rem;">🧩</div><div><p style="margin:0;color:#004E92;font-size:1.5rem;font-weight:800;">Adaptador V11.0: Profissional</p></div></div>""", unsafe_allow_html=True)

if not st.session_state.banco_estudantes:
    st.warning("⚠️ Cadastre um aluno no PEI 360º primeiro.")
    st.stop()

lista = [a['nome'] for a in st.session_state.banco_estudantes]
nome_aluno = st.selectbox("📂 Selecione o Estudante:", lista)
aluno = next(a for a in st.session_state.banco_estudantes if a['nome'] == nome_aluno)

st.markdown(f"""
    <div class="student-header">
        <div class="student-info-item"><div class="student-label">Nome</div><div class="student-value">{aluno.get('nome')}</div></div>
        <div class="student-info-item"><div class="student-label">Série</div><div class="student-value">{aluno.get('serie', '-')}</div></div>
        <div class="student-info-item"><div class="student-label">Hiperfoco</div><div class="student-value">{aluno.get('hiperfoco', '-')}</div></div>
    </div>
""", unsafe_allow_html=True)

# 5 ABAS REESTRUTURADAS
tab_docx, tab_img, tab_create, tab_visual, tab_ctx = st.tabs([
    "📄 Adaptar Prova (Word)", 
    "✂️ Adaptar Atividade (Print)", 
    "✨ Criar do Zero", 
    "🎨 Estúdio Visual", 
    "💡 Contextualizar"
])

# 1. ADAPTAR PROVA (WORD)
with tab_docx:
    c1, c2, c3 = st.columns(3)
    materia_d = c1.selectbox("Matéria", ["Matemática", "Português", "Ciências", "História", "Geografia", "Artes", "Ed. Física", "Inglês"], key="dm")
    tema_d = c2.text_input("Tema", placeholder="Ex: Frações", key="dt")
    tipo_d = c3.selectbox("Tipo", ["Prova", "Tarefa"], key="dtp")
    arquivo_d = st.file_uploader("Upload DOCX", type=["docx"], key="fd")
    
    if 'docx_imgs' not in st.session_state: st.session_state.docx_imgs = []
    if 'docx_txt' not in st.session_state: st.session_state.docx_txt = None
    
    if arquivo_d and arquivo_d.file_id != st.session_state.get('last_d'):
        st.session_state.last_d = arquivo_d.file_id
        txt, imgs = extrair_dados_docx(arquivo_d)
        st.session_state.docx_txt = txt; st.session_state.docx_imgs = imgs
        st.success(f"{len(imgs)} imagens.")

    map_d = {}; qs_d = []
    if st.session_state.docx_imgs:
        st.write("### Mapeamento")
        cols = st.columns(3)
        for i, img in enumerate(st.session_state.docx_imgs):
            with cols[i % 3]:
                st.image(img, width=80)
                q = st.number_input(f"Questão:", 0, 50, key=f"dq_{i}")
                if q > 0: map_d[int(q)] = img; qs_d.append(int(q))

    if st.button("🚀 ADAPTAR PROVA", type="primary", key="btn_d"):
        if not st.session_state.docx_txt: st.warning("Envie arquivo."); st.stop()
        with st.spinner("Adaptando com inteligência PEI..."):
            rac, txt = adaptar_conteudo_docx(api_key, aluno, st.session_state.docx_txt, materia_d, tema_d, tipo_d, True, qs_d)
            st.session_state['res_docx'] = {'rac': rac, 'txt': txt, 'map': map_d}
            st.rerun()

    if 'res_docx' in st.session_state:
        res = st.session_state['res_docx']
        with st.expander("🧠 Racional Pedagógico"): st.info(res['rac'])
        with st.container(border=True):
            partes = re.split(r'(\[\[IMG_Q\d+\]\])', res['txt'])
            for p in partes:
                tag = re.search(r'\[\[IMG_Q(\d+)\]\]', p)
                if tag:
                    i = int(tag.group(1))
                    im = res['map'].get(i)
                    if im: st.image(im, width=300)
                elif p.strip(): st.markdown(p.strip())
        docx = construir_docx_final(res['txt'], aluno, materia_d, res['map'], None, tipo_d)
        st.download_button("📥 BAIXAR DOCX", docx, "Prova.docx", "primary")

# 2. ADAPTAR ATIVIDADE (PRINT)
with tab_img:
    st.info("Adaptar a partir de Imagem/Foto.")
    c1, c2, c3 = st.columns(3)
    # Lista Ampliada
    discip = ["Matemática", "Português", "Ciências", "História", "Geografia", "Artes", "Ed. Física", "Inglês", "Filosofia", "Sociologia", "Física", "Química", "Biologia"]
    materia_i = c1.selectbox("Matéria", discip, key="im")
    tema_i = c2.text_input("Tema", placeholder="Ex: Células", key="it")
    tipo_i = c3.selectbox("Tipo", ["Atividade", "Tarefa"], key="itp")
    arquivo_i = st.file_uploader("Upload Imagem", type=["png","jpg","jpeg"], key="fi")
    
    if 'img_raw' not in st.session_state: st.session_state.img_raw = None
    if arquivo_i and arquivo_i.file_id != st.session_state.get('last_i'):
        st.session_state.last_i = arquivo_i.file_id
        img = Image.open(arquivo_i).convert("RGB")
        buf = BytesIO(); img.save(buf, format="JPEG"); st.session_state.img_raw = buf.getvalue()

    cropped_res = None
    if st.session_state.img_raw:
        st.markdown("### ✂️ Recorte e Gere")
        img_pil = Image.open(BytesIO(st.session_state.img_raw))
        img_pil.thumbnail((800, 800))
        # Realtime=True para feedback visual, mas processamento só no botão
        cropped_res = st_cropper(img_pil, realtime_update=True, box_color='#FF0000', aspect_ratio=None, key="crop_i")
        if cropped_res: st.image(cropped_res, width=200, caption="Prévia")

    if st.button("🚀 ADAPTAR RECORTE", type="primary", key="btn_i"):
        if not cropped_res: st.warning("Recorte inválido."); st.stop()
        with st.spinner("Lendo texto e adaptando..."):
            buf_c = BytesIO()
            cropped_res.convert('RGB').save(buf_c, format="JPEG", quality=85)
            img_bytes = buf_c.getvalue()
            
            rac, txt = adaptar_conteudo_imagem(api_key, aluno, img_bytes, materia_i, tema_i, tipo_i, True)
            st.session_state['res_img'] = {'rac': rac, 'txt': txt, 'map': {1: img_bytes}}
            st.rerun()

    if 'res_img' in st.session_state:
        res = st.session_state['res_img']
        with st.expander("🧠 Racional"): st.info(res['rac'])
        with st.container(border=True):
            partes = re.split(r'(\[\[IMG_Q\d+\]\])', res['txt'])
            for p in partes:
                tag = re.search(r'\[\[IMG_Q(\d+)\]\]', p)
                if tag:
                    im = res['map'].get(1)
                    if im: st.image(im, width=300)
                elif p.strip(): st.markdown(p.strip())
        docx = construir_docx_final(res['txt'], aluno, materia_i, res['map'], None, tipo_i)
        st.download_button("📥 BAIXAR DOCX", docx, "Atividade.docx", "primary")

# 3. CRIAR DO ZERO (PROFISSIONAL)
with tab_create:
    st.info("Criação de Alto Nível (Professor Sênior).")
    cc1, cc2 = st.columns(2)
    mat_c = cc1.selectbox("Componente", discip, key="cm")
    obj_c = cc2.text_input("Assunto/Objeto de Conhecimento", key="co")
    qtd_c = st.slider("Qtd Questões", 1, 10, 5, key="cq")
    
    if st.button("✨ CRIAR ATIVIDADE PROFISSIONAL", type="primary", key="btn_c"):
        with st.spinner("Elaborando questões densas e contextualizadas..."):
            rac, txt = criar_profissional(api_key, aluno, mat_c, obj_c, qtd_c, "Mista")
            
            # Geração de Imagens
            novo_map = {}; count = 0
            tags = re.findall(r'\[\[GEN_IMG: (.*?)\]\]', txt)
            for p in tags:
                count += 1
                url = gerar_dalle_prompt(api_key, p)
                if url:
                    io = baixar_imagem_url(url)
                    if io: novo_map[count] = io.getvalue()
            
            txt_fin = txt
            for i in range(1, count + 1): 
                txt_fin = re.sub(r'\[\[GEN_IMG: .*?\]\]', f"[[IMG_G{i}]]", txt_fin, count=1)
                
            st.session_state['res_create'] = {'rac': rac, 'txt': txt_fin, 'map': novo_map}
            st.rerun()

    if 'res_create' in st.session_state:
        res = st.session_state['res_create']
        with st.expander("🧠 Racional"): st.info(res['rac'])
        with st.container(border=True):
            partes = re.split(r'(\[\[IMG_G\d+\]\])', res['txt'])
            for p in partes:
                tag = re.search(r'\[\[IMG_G(\d+)\]\]', p)
                if tag:
                    i = int(tag.group(1))
                    im = res['map'].get(i)
                    if im: st.image(im, width=300)
                elif p.strip(): st.markdown(p.strip())
        docx = construir_docx_final(res['txt'], aluno, mat_c, {}, None, "Criada")
        st.download_button("📥 BAIXAR DOCX", docx, "Criada.docx", "primary")

# 4. ESTÚDIO VISUAL (MÁGICA DO PEI)
with tab_visual:
    c_a, c_b = st.columns([3, 1])
    desc = c_a.text_area("Descrição Manual:", placeholder="Ex: Rotina de banheiro...", key="vd")
    
    if c_b.button("✨ Mágica do PEI", help="Gera uma imagem baseada no que o aluno precisa segundo o PEI"):
        with st.spinner("Lendo PEI e imaginando..."):
            desc_auto = sugerir_imagem_pei(api_key, aluno)
            st.session_state['img_prompt'] = desc_auto
            st.rerun()
            
    if 'img_prompt' in st.session_state:
        st.info(f"Sugestão Automática: {st.session_state['img_prompt']}")
        if st.button("🎨 Gerar Sugestão"):
             with st.spinner("Desenhando..."):
                url = gerar_dalle_prompt(api_key, st.session_state['img_prompt'])
                if url: st.image(url)

    if st.button("🎨 Gerar Manual"):
        with st.spinner("Desenhando..."):
            url = gerar_dalle_prompt(api_key, f"{desc} style {aluno.get('hiperfoco')}")
            if url: st.image(url)

# 5. CONTEXTUALIZAR
with tab_ctx:
    st.info("Roteiro de Aula Rico.")
    c1, c2 = st.columns(2)
    ass = c1.text_input("O que você vai ensinar?", placeholder="Ex: Fotossíntese", key="cxa")
    t_extra = c2.text_input("Tema de Conexão (Opcional)", value=aluno.get('hiperfoco'), key="cxb")
    
    if st.button("💡 CRIAR ROTEIRO", type="primary"):
        with st.spinner("Planejando..."):
            st.markdown(gerar_contextualizacao(api_key, aluno, ass, t_extra))
